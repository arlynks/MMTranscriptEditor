import customtkinter as ctk
from tkinter import filedialog
from faster_whisper import WhisperModel
import torch
import os
import json
import re
import threading
import platform
import subprocess
import shutil
import tempfile
import time
import sys
import urllib.request
import urllib.error
from PIL import Image

# Simple version comparison function
def compare_versions(current, latest):
    """
    Compare two version strings.
    Returns: 1 if latest > current, 0 if equal, -1 if latest < current
    """
    try:
        # Remove 'v' prefix if present
        current = current.lstrip("vV")
        latest = latest.lstrip("vV")
        
        # Split into parts
        current_parts = [int(x) for x in current.split(".")]
        latest_parts = [int(x) for x in latest.split(".")]
        
        # Pad with zeros to match length
        max_len = max(len(current_parts), len(latest_parts))
        current_parts.extend([0] * (max_len - len(current_parts)))
        latest_parts.extend([0] * (max_len - len(latest_parts)))
        
        # Compare
        for c, l in zip(current_parts, latest_parts):
            if l > c:
                return 1
            elif l < c:
                return -1
        return 0
    except Exception:
        # If comparison fails, assume versions are equal
        return 0
# Transformers pipeline will be lazy-loaded to avoid startup issues
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

# Try to import pygame for audio playback
PYGAME_AVAILABLE = False
try:
    import pygame
    # Try to initialize mixer with error handling
    try:
        pygame.mixer.pre_init(frequency=44100, size=-16, channels=2, buffer=512)
        pygame.mixer.init()
        # Test if mixer is actually working
        if pygame.mixer.get_init():
            PYGAME_AVAILABLE = True
        else:
            print("Warning: pygame mixer initialized but not working properly")
            PYGAME_AVAILABLE = False
    except Exception as init_error:
        print(f"Warning: pygame mixer initialization failed: {init_error}")
        # Try default initialization as fallback
        try:
            pygame.mixer.init()
            if pygame.mixer.get_init():
                PYGAME_AVAILABLE = True
        except:
            PYGAME_AVAILABLE = False
except ImportError:
    PYGAME_AVAILABLE = False
    print("Warning: pygame not available. Audio player will not work.")
except Exception as e:
    PYGAME_AVAILABLE = False
    print(f"Warning: pygame import failed: {e}")

# Try to import pydub for audio speed control
try:
    from pydub import AudioSegment
    from pydub.playback import play
    PYDUB_AVAILABLE = True
except ImportError:
    PYDUB_AVAILABLE = False
    print("Warning: pydub not available. Audio speed control will be limited.")
    print("To install pydub, run in your terminal:")
    print("  grammateus-env\\Scripts\\activate")
    print("  pip install pydub")
    print("  pip install ffmpeg-python  # if ffmpeg not on PATH")
except Exception as e:
    PYDUB_AVAILABLE = False
    print(f"Warning: pydub initialization failed: {e}")

# Try to import mutagen for audio metadata
try:
    from mutagen import File as MutagenFile
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False

# Import license management
try:
    from mmvidstoclips_license import (
        initialize_trial,
        get_license_status,
        is_unlocked,
        is_trial_expired,
        can_run,
        unlock_with_key,
        get_trial_remaining
    )
    LICENSE_AVAILABLE = True
except ImportError as e:
    LICENSE_AVAILABLE = False
    print(f"Warning: License module not available: {e}")
    print("Application will run without license checks.")
except Exception as e:
    LICENSE_AVAILABLE = False
    print(f"Warning: Error loading license module: {e}")
    print("Application will run without license checks.")

# --- Application Version ---
CURRENT_VERSION = "1.0.0"

# --- Update Check Function ---
def check_for_update():
    """
    Check for updates by calling GitHub Releases API.
    Returns a tuple: (success: bool, message: str)
    """
    try:
        # GitHub Releases API endpoint
        api_url = "https://api.github.com/repos/arlynks/MMTranscriptEditor/releases/latest"
        
        # Make request with timeout
        request = urllib.request.Request(api_url)
        request.add_header('User-Agent', 'MMTranscriptEditor-UpdateChecker/1.0')
        
        with urllib.request.urlopen(request, timeout=10) as response:
            if response.status != 200:
                return False, "Unable to check for updates right now. Please try again later."
            
            data = json.loads(response.read().decode())
            
            # Get the latest version tag (remove 'v' prefix if present)
            latest_tag = data.get("tag_name", "")
            latest_version = latest_tag.lstrip("vV") if latest_tag else ""
            
            # Compare versions
            try:
                if latest_version and compare_versions(CURRENT_VERSION, latest_version) > 0:
                    # New version available - get download link
                    assets = data.get("assets", [])
                    download_url = ""
                    if assets:
                        download_url = assets[0].get("browser_download_url", "")
                    
                    message = f"A new version is available!\n\n"
                    message += f"Current version: {CURRENT_VERSION}\n"
                    message += f"Latest version: {latest_version}\n\n"
                    
                    if download_url:
                        message += f"Download: {download_url}\n\n"
                    else:
                        message += f"Visit: https://github.com/arlynks/MMTranscriptEditor/releases/latest\n\n"
                    
                    message += "Would you like to open the download page?"
                    return True, message
                else:
                    return True, f"You're on the latest version ({CURRENT_VERSION})."
            except Exception as e:
                # Version comparison failed - assume up to date
                return True, f"You're on the latest version ({CURRENT_VERSION})."
    
    except urllib.error.URLError as e:
        # Network error (no internet, timeout, etc.)
        return False, "Unable to check for updates right now. Please try again later."
    except json.JSONDecodeError:
        # Invalid JSON response
        return False, "Unable to check for updates right now. Please try again later."
    except Exception as e:
        # Any other error
        return False, "Unable to check for updates right now. Please try again later."

# --- Settings Management ---
# Handle PyInstaller bundled executables
if getattr(sys, 'frozen', False):
    # Running as compiled executable
    if hasattr(sys, '_MEIPASS'):
        # PyInstaller onefile mode - use user's AppData for data files
        APP_DIR = os.path.join(os.path.expanduser("~"), "MMTranscriptEditor")
    else:
        # PyInstaller directory mode - use executable directory
        APP_DIR = os.path.dirname(os.path.abspath(sys.executable))
else:
    # Running as script
    APP_DIR = os.path.dirname(os.path.abspath(__file__))

# Ensure APP_DIR exists
if not os.path.exists(APP_DIR):
    try:
        os.makedirs(APP_DIR, exist_ok=True)
    except:
        # Fallback to current directory if we can't create APP_DIR
        APP_DIR = os.getcwd()

SETTINGS_FILE = os.path.join(APP_DIR, "mmvidstoclips_settings.json")
SAVED_TRANSCRIPTS_DIR = os.path.join(APP_DIR, "saved_transcripts")

# Create saved transcripts folder if it doesn't exist
if not os.path.exists(SAVED_TRANSCRIPTS_DIR):
    os.makedirs(SAVED_TRANSCRIPTS_DIR)

DEFAULT_SETTINGS = {
    "theme": "Dark",
    "whisper_model": "small",
    "window_geometry": "1560x980+100+100"  # Default: width x height + x + y
}

def load_settings():
    """Load user settings from file."""
    try:
        if os.path.exists(SETTINGS_FILE):
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                settings = json.load(f)
                # Merge with defaults to handle missing keys
                return {**DEFAULT_SETTINGS, **settings}
    except Exception:
        pass
    return DEFAULT_SETTINGS.copy()

def save_settings(settings):
    """Save user settings to file."""
    try:
        with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
            json.dump(settings, f, indent=2)
    except Exception as e:
        print(f"Error saving settings: {e}")

# --- Saved Transcripts Management ---
def get_saved_transcripts():
    """Get list of saved transcript filenames."""
    if not os.path.exists(SAVED_TRANSCRIPTS_DIR):
        return []
    files = [f for f in os.listdir(SAVED_TRANSCRIPTS_DIR) if f.endswith('.txt')]
    return sorted(files, key=lambda x: os.path.getmtime(os.path.join(SAVED_TRANSCRIPTS_DIR, x)), reverse=True)

def save_transcript_to_app(text, filename=None, formatting_tags=None):
    """Save transcript text to the in-app saved_transcripts folder.
    
    Args:
        text: The text content to save
        filename: Optional filename (will generate if None)
        formatting_tags: Optional dict of formatting tags to save
    """
    if not text.strip():
        return None
    
    if filename is None:
        # Generate filename with timestamp
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"transcript_{timestamp}.txt"
    
    if not filename.endswith('.txt'):
        filename += '.txt'
    
    filepath = os.path.join(SAVED_TRANSCRIPTS_DIR, filename)
    try:
        # Save text file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(text)
        
        # Save formatting tags if provided
        if formatting_tags:
            format_filepath = filepath.replace('.txt', '.formatting.json')
            with open(format_filepath, 'w', encoding='utf-8') as f:
                json.dump(formatting_tags, f, indent=2)
        
        return filename
    except Exception as e:
        print(f"Error saving transcript: {e}")
        return None

def load_transcript_from_app(filename):
    """Load transcript text and formatting from the in-app saved_transcripts folder.
    
    Returns:
        tuple: (text_content, formatting_tags_dict) or (text_content, None) if no formatting
    """
    filepath = os.path.join(SAVED_TRANSCRIPTS_DIR, filename)
    try:
        # Load text file
        with open(filepath, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Try to load formatting tags
        format_filepath = filepath.replace('.txt', '.formatting.json')
        formatting_tags = None
        if os.path.exists(format_filepath):
            try:
                with open(format_filepath, 'r', encoding='utf-8') as f:
                    formatting_tags = json.load(f)
            except Exception as e:
                print(f"Error loading formatting tags: {e}")
        
        return (text, formatting_tags) if formatting_tags else (text, None)
    except Exception as e:
        print(f"Error loading transcript: {e}")
        return (None, None)

def delete_saved_transcript(filename):
    """Delete a saved transcript from the in-app folder."""
    filepath = os.path.join(SAVED_TRANSCRIPTS_DIR, filename)
    try:
        if os.path.exists(filepath):
            os.remove(filepath)
            return True
    except Exception as e:
        print(f"Error deleting transcript: {e}")
    return False

def refresh_saved_transcripts_dropdown():
    """Refresh the saved transcripts dropdown with current files."""
    files = get_saved_transcripts()
    if files:
        saved_transcripts_dropdown.configure(values=["-- Select --"] + files)
    else:
        saved_transcripts_dropdown.configure(values=["(No saved transcripts)"])

# Load settings on startup
user_settings = load_settings()

# --- Theme Configuration ---
CURRENT_THEME = user_settings.get("theme", "Dark")
ctk.set_appearance_mode(CURRENT_THEME)
ctk.set_default_color_theme("blue")

# --- Theme Colors ---
THEME_COLORS = {
    "Dark": {
        "textbox_bg": "#666666",  # 2 shades lighter grey for Dark theme
        "textbox_border": "#505050",
        "textbox_text": "white",
        "resize_handle": "#505050",
        "mini_bar": "#2D2D2D",
        "main_bg": "#1E1E1E"  # Dark grey for main UI
    },
    "Light": {
        "textbox_bg": "#F5F5F5",  # Slightly lighter grey (was white)
        "textbox_border": "#CCCCCC",
        "textbox_text": "black",
        "resize_handle": "#CCCCCC",
        "mini_bar": "#E0E0E0",
        "main_bg": "#D0D0D0"  # Slightly darker grey for main UI background (header and sidebar)
    }
}

def get_theme_color(key):
    """Get color for current theme."""
    return THEME_COLORS.get(CURRENT_THEME, THEME_COLORS["Dark"]).get(key)

# --- Button Color Configuration ---
# Medium grey buttons with white text, darker grey on hover
BUTTON_COLOR = "#808080"  # Medium grey
BUTTON_HOVER_COLOR = "#606060"  # Darker grey on hover
BUTTON_TEXT_COLOR = "white"  # White text

# --- Modern Font Configuration ---
FONT_FAMILY = "Segoe UI"  # Modern sans-serif (fallback: system default)
FONT_SIZES = {
    "title": 18,
    "heading": 15,
    "body": 13,
    "small": 11,
    "button": 13
}

# --- Device Detection ---
USE_CUDA = torch.cuda.is_available()
DEVICE = "cuda" if USE_CUDA else "cpu"
COMPUTE_TYPE = "float16" if USE_CUDA else "int8"  # int8 is faster on CPU

# Available Whisper models (faster-whisper supports all on CPU efficiently)
AVAILABLE_MODELS = ["tiny", "base", "small", "medium", "large-v2"]
DEFAULT_MODEL = user_settings.get("whisper_model", "small")

# Load faster-whisper model
model = WhisperModel(DEFAULT_MODEL, device=DEVICE, compute_type=COMPUTE_TYPE)
current_model_name = DEFAULT_MODEL

# Lazy-load summarizer to avoid startup import issues with PyInstaller
_summarizer = None

def get_summarizer():
    """Lazy-load the transformers pipeline only when needed."""
    global _summarizer
    if _summarizer is None:
        try:
            from transformers import pipeline
            _summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
        except Exception as e:
            print(f"Error loading summarizer: {e}")
            return None
    return _summarizer

# --- Main App Window ---
app = ctk.CTk()
app.title("MM Transcript Editor")

# Set application icon
ICON_PATH = None
# Try .ico file first (Windows)
ico_path = os.path.join(APP_DIR, "MMTranscriptEditor.ico")
if os.path.exists(ico_path):
    ICON_PATH = ico_path
    try:
        app.iconbitmap(ico_path)
    except:
        pass
else:
    # Fallback to thumb.png
    thumb_path = os.path.join(APP_DIR, "thumb.png")
    if os.path.exists(thumb_path):
        ICON_PATH = thumb_path
        try:
            # Convert PNG to icon for Windows
            icon_image = Image.open(thumb_path)
            # Save as temporary ICO if needed
            temp_ico = os.path.join(APP_DIR, "temp_icon.ico")
            try:
                icon_image.save(temp_ico, format='ICO')
                app.iconbitmap(temp_ico)
            except:
                # Try using iconphoto as fallback
                from PIL import ImageTk
                photo = ImageTk.PhotoImage(icon_image)
                app.iconphoto(False, photo)
        except Exception as e:
            print(f"Could not set application icon: {e}")

# Load saved window geometry or use default
saved_geometry = user_settings.get("window_geometry", "1560x980+100+100")
app.geometry(saved_geometry)

# Function to save window geometry
def save_window_geometry():
    """Save current window geometry to settings."""
    try:
        geometry = app.geometry()
        user_settings["window_geometry"] = geometry
        save_settings(user_settings)
    except Exception as e:
        print(f"Error saving window geometry: {e}")

# Save geometry when window is closed
def on_closing():
    """Handle window closing event."""
    save_window_geometry()
    app.destroy()

app.protocol("WM_DELETE_WINDOW", on_closing)

# Save geometry when window is moved or resized (with debouncing)
_geometry_save_timer = None
_last_geometry = app.geometry()

def on_window_configure(event):
    """Handle window configuration changes (move/resize)."""
    global _last_geometry, _geometry_save_timer
    # Only track changes to the main window, not child widgets
    if event.widget == app:
        current_geometry = app.geometry()
        # Only save if geometry actually changed
        if current_geometry != _last_geometry:
            _last_geometry = current_geometry
            # Debounce: cancel previous timer and schedule new save
            if _geometry_save_timer:
                app.after_cancel(_geometry_save_timer)
            _geometry_save_timer = app.after(500, save_window_geometry)  # Save 500ms after last change

# Bind to window configuration changes
app.bind("<Configure>", on_window_configure)

# Global keyboard shortcuts
def on_global_key_press(event):
    """Handle global keyboard shortcuts."""
    # Ctrl+F for search - find focused panel and search
    if event.state & 0x4 and event.keysym.lower() == 'f':  # Ctrl+F
        # Find the panel with focus
        for panel in panels:
            try:
                if panel.textbox._textbox.focus_get() == panel.textbox._textbox:
                    panel.show_search_dialog()
                    return "break"
            except:
                pass

app.bind("<Control-f>", on_global_key_press)
app.bind("<Control-F>", on_global_key_press)

# Apply theme-specific background color
def apply_main_bg_color():
    """Apply main background color based on theme."""
    bg_color = get_theme_color("main_bg")
    if bg_color:
        app.configure(fg_color=bg_color)

apply_main_bg_color()

# --- Export Functions ---
def export_to_txt(text, file_path):
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(text)

def extract_formatted_text_segments(textbox_widget):
    """Extract text with formatting information from a tkinter textbox.
    Preserves newline structure to detect paragraph breaks."""
    segments = []
    content = textbox_widget.get("1.0", "end-1c")  # Get all text
    
    if not content:
        return segments
    
    # Iterate character by character through the entire content to preserve newlines
    current_tags = None
    current_text = ""
    line_num = 1
    char_idx = 0
    
    for i, char in enumerate(content):
        # Calculate tkinter position (line.char format)
        # Count newlines before this character
        newlines_before = content[:i].count('\n')
        line_num = newlines_before + 1
        # Character index within current line
        if newlines_before > 0:
            last_newline_pos = content[:i].rfind('\n')
            char_idx = i - last_newline_pos - 1
        else:
            char_idx = i
        
        pos = f"{line_num}.{char_idx}"
        
        try:
            # Get all tags at this position
            tags = list(textbox_widget.tag_names(pos))
            # Filter out internal tags
            tags = [t for t in tags if t not in ("sel", "insert", "current")]
        except:
            tags = []
        
        # Normalize tags (sort for comparison)
        tags_key = tuple(sorted(tags))
        
        # Check if formatting changed (new tags or tag removed)
        if tags_key != current_tags:
            # Save previous segment
            if current_text:
                segments.append({
                    'text': current_text,
                    'tags': list(current_tags) if current_tags else []
                })
            
            # Start new segment with current character
            current_tags = tags_key
            current_text = char
        else:
            # Same formatting - append to current segment
            current_text += char
    
    # Add last segment
    if current_text:
        segments.append({
            'text': current_text,
            'tags': list(current_tags) if current_tags else []
        })
    
    return segments

def apply_docx_formatting(run, tags, textbox, formatting_tags_dict=None):
    """Apply formatting to a DOCX run based on tkinter textbox tags."""
    try:
        if not tags:
            return
        
        # Check each tag for formatting properties
        for tag in tags:
            try:
                # First check formatting_tags_dict if provided (more reliable)
                if formatting_tags_dict and tag in formatting_tags_dict:
                    tag_info = formatting_tags_dict[tag]
                    tag_type = tag_info.get("type")
                    
                    # Handle italic tag
                    if tag_type == "italic" or tag_info.get("italic", False):
                        run.italic = True
                    
                    # Handle bold tag
                    if tag_type == "bold" or tag_info.get("bold", False):
                        run.bold = True
                    
                    # Handle font_merged tag
                    if tag_type == "font_merged":
                        if tag_info.get("italic", False):
                            run.italic = True
                        if tag_info.get("bold", False):
                            run.bold = True
                        if tag_info.get("size"):
                            try:
                                run.font.size = Pt(int(tag_info.get("size")))
                            except:
                                pass
                        if tag_info.get("family"):
                            try:
                                run.font.name = str(tag_info.get("family"))
                            except:
                                pass
                    
                    # Handle highlight tag
                    if tag_type == "highlight":
                        highlight_color = tag_info.get("color")
                        if highlight_color:
                            try:
                                # DOCX has limited highlight colors, use yellow as default
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            except:
                                pass
                
                # Fall back to tkinter tag config
                tag_config = textbox.tag_config(tag)
                
                # Check for standalone italic tag name
                if tag == "italic" or "italic" in tag.lower():
                    run.italic = True
                
                # Check for standalone bold tag name
                if tag == "bold" or "bold" in tag.lower():
                    run.bold = True
                
                # Bold, italic, font
                if 'font' in tag_config:
                    font_value = tag_config['font'][4] if len(tag_config['font']) > 4 else None
                    if font_value:
                        if isinstance(font_value, tuple) and len(font_value) > 2:
                            if 'bold' in font_value[2:]:
                                run.bold = True
                            if 'italic' in font_value[2:]:
                                run.italic = True
                            # Font size
                            if len(font_value) > 1:
                                try:
                                    size = int(font_value[1])
                                    run.font.size = Pt(size)
                                except:
                                    pass
                            # Font family
                            if len(font_value) > 0:
                                try:
                                    run.font.name = str(font_value[0])
                                except:
                                    pass
                
                # Underline
                if 'underline' in tag_config:
                    underline_value = tag_config['underline'][4] if len(tag_config['underline']) > 4 else False
                    if underline_value:
                        run.underline = True
                
                # Foreground color
                if 'foreground' in tag_config:
                    fg_color = tag_config['foreground'][4] if len(tag_config['foreground']) > 4 else None
                    if fg_color:
                        try:
                            hex_color = fg_color.lstrip('#')
                            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                            run.font.color.rgb = RGBColor(*rgb)
                        except:
                            pass
                
                # Background/highlight (limited support in DOCX)
                if 'background' in tag_config:
                    bg_color = tag_config['background'][4] if len(tag_config['background']) > 4 else None
                    if bg_color:
                        try:
                            # DOCX has limited highlight colors, use yellow as default
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        except:
                            pass
            except Exception:
                # Tag might not exist or be configured
                pass
    except Exception as e:
        print(f"Error applying DOCX formatting: {e}")

def map_font_to_reportlab(font_family, bold=False, italic=False):
    """Map common font families to ReportLab's built-in fonts.
    ReportLab supports: Helvetica, Times-Roman, Courier (and their variants)."""
    # Map common fonts to ReportLab built-in fonts
    font_map = {
        'Segoe UI': 'Helvetica',
        'Arial': 'Helvetica',
        'Calibri': 'Helvetica',
        'Verdana': 'Helvetica',
        'Tahoma': 'Helvetica',
        'Trebuchet MS': 'Helvetica',
        'Times New Roman': 'Times-Roman',
        'Times': 'Times-Roman',
        'Georgia': 'Times-Roman',
        'Courier New': 'Courier',
        'Courier': 'Courier',
        'Consolas': 'Courier',
        'Lucida Console': 'Courier',
        'Monaco': 'Courier',
    }
    
    # Get base font name (default to Helvetica if not mapped)
    base_font = font_map.get(font_family, 'Helvetica')
    
    # Build variant name based on bold/italic
    if bold and italic:
        if base_font == 'Helvetica':
            return 'Helvetica-BoldOblique'
        elif base_font == 'Times-Roman':
            return 'Times-BoldItalic'
        elif base_font == 'Courier':
            return 'Courier-BoldOblique'
    elif bold:
        if base_font == 'Helvetica':
            return 'Helvetica-Bold'
        elif base_font == 'Times-Roman':
            return 'Times-Bold'
        elif base_font == 'Courier':
            return 'Courier-Bold'
    elif italic:
        if base_font == 'Helvetica':
            return 'Helvetica-Oblique'
        elif base_font == 'Times-Roman':
            return 'Times-Italic'
        elif base_font == 'Courier':
            return 'Courier-Oblique'
    
    return base_font

def apply_pdf_formatting_styles(text, tags, textbox, formatting_tags_dict=None):
    """Build styled XML-like text for PDF from tkinter textbox tags.
    ReportLab uses XML-like tags with specific syntax."""
    try:
        if not tags:
            # Still escape XML characters
            return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Escape XML special characters in text
        text_escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        style_parts = []
        close_tags = []
        font_attrs = {}
        is_bold = False
        is_italic = False
        back_color = None
        
        # Check each tag for formatting properties
        for tag in tags:
            try:
                # First check formatting_tags dict if provided (more reliable)
                if formatting_tags_dict and tag in formatting_tags_dict:
                    tag_info = formatting_tags_dict[tag]
                    tag_type = tag_info.get("type")
                    
                    if tag_type == "font_merged":
                        # Apply from formatting_tags dict
                        is_bold = tag_info.get("bold", False)
                        is_italic = tag_info.get("italic", False)
                        size_val = tag_info.get("size")
                        if size_val:
                            try:
                                font_attrs['fontSize'] = str(int(size_val))
                            except:
                                pass
                        family = tag_info.get("family", "Helvetica")
                        # Map font to ReportLab built-in and handle bold/italic in font name
                        mapped_font = map_font_to_reportlab(family, is_bold, is_italic)
                        font_attrs['fontName'] = mapped_font
                        continue
                    elif tag_type == "italic":
                        is_italic = True
                    elif tag_type == "bold":
                        is_bold = True
                    elif tag_type == "highlight":
                        # Get highlight color
                        highlight_color = tag_info.get("color")
                        if highlight_color:
                            back_color = highlight_color
                        continue
                
                # Fall back to tkinter tag config
                tag_config = textbox.tag_config(tag)
                
                # Check for standalone italic/bold tag names
                if tag == "italic" or "italic" in tag.lower():
                    is_italic = True
                if tag == "bold" or "bold" in tag.lower():
                    is_bold = True
                
                # Font formatting
                if 'font' in tag_config:
                    font_value = tag_config['font'][4] if len(tag_config['font']) > 4 else None
                    if font_value and isinstance(font_value, tuple):
                        if len(font_value) > 2:
                            if 'bold' in font_value[2:]:
                                is_bold = True
                            if 'italic' in font_value[2:]:
                                is_italic = True
                        # Font size
                        if len(font_value) > 1:
                            try:
                                font_attrs['fontSize'] = str(int(font_value[1]))
                            except:
                                pass
                        # Font family - map to ReportLab built-in
                        if len(font_value) > 0:
                            try:
                                family = str(font_value[0])
                                mapped_font = map_font_to_reportlab(family, is_bold, is_italic)
                                font_attrs['fontName'] = mapped_font
                            except:
                                pass
                
                # Underline
                if 'underline' in tag_config:
                    underline_value = tag_config['underline'][4] if len(tag_config['underline']) > 4 else False
                    if underline_value:
                        style_parts.append('<u>')
                        close_tags.insert(0, '</u>')
                
                # Foreground color
                if 'foreground' in tag_config:
                    fg_color = tag_config['foreground'][4] if len(tag_config['foreground']) > 4 else None
                    if fg_color:
                        font_attrs['textColor'] = fg_color
                
                # Background/highlight color
                if 'background' in tag_config:
                    bg_color = tag_config['background'][4] if len(tag_config['background']) > 4 else None
                    if bg_color:
                        back_color = bg_color
            except Exception:
                pass
        
        # Apply bold/italic to font if not already set
        if is_bold or is_italic:
            if 'fontName' not in font_attrs:
                # Set default font with bold/italic
                font_attrs['fontName'] = map_font_to_reportlab("Helvetica", is_bold, is_italic)
            else:
                # Update existing font name to include bold/italic
                current_font = font_attrs.get('fontName', 'Helvetica')
                font_attrs['fontName'] = map_font_to_reportlab(current_font.split('-')[0] if '-' in current_font else current_font, is_bold, is_italic)
        
        # Build font tag if we have font attributes
        # ReportLab expects fontSize and fontName (camelCase), not size and name
        if font_attrs:
            # Build attribute string with proper quoting
            attr_parts = []
            for k, v in font_attrs.items():
                # Remove any existing quotes and add proper ones
                v_clean = v.strip('"\'')
                # Escape quotes in the value
                v_clean = v_clean.replace('"', '&quot;')
                # Always quote fontName and fontSize
                attr_parts.append(f'{k}="{v_clean}"')
            font_attr_str = ' '.join(attr_parts)
            style_parts.append(f'<font {font_attr_str}>')
            close_tags.insert(0, '</font>')
        
        # Add background color if present (ReportLab uses backColor attribute)
        if back_color:
            try:
                # Ensure color is in hex format
                if not back_color.startswith('#'):
                    back_color = '#' + back_color.lstrip('#')
                # Escape quotes
                back_color_clean = back_color.replace('"', '&quot;')
                style_parts.append(f'<font backColor="{back_color_clean}">')
                close_tags.insert(0, '</font>')
            except Exception as e:
                print(f"Error adding background color: {e}")
        
        # Combine opening tags, text, and closing tags
        result = ''.join(style_parts) + text_escaped + ''.join(close_tags)
        return result
    except Exception as e:
        print(f"Error in apply_pdf_formatting_styles: {e}")
        return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

def export_to_docx(text, file_path, formatting_tags=None, textbox_widget=None):
    """Export text to DOCX with formatting support."""
    doc = Document()
    
    # Use textbox widget if available for better formatting extraction
    if textbox_widget and hasattr(textbox_widget, '_textbox'):
        try:
            textbox = textbox_widget._textbox
            segments = extract_formatted_text_segments(textbox)
            
            # Split full text into paragraphs by double newlines first
            full_text = textbox.get("1.0", "end-1c")
            paragraphs_raw = full_text.split('\n\n')
            
            # Process each paragraph
            char_offset = 0
            for para_idx, para_text in enumerate(paragraphs_raw):
                # Clean paragraph text (replace single newlines with spaces)
                para_text_clean = para_text.replace('\n', ' ').strip()
                
                if not para_text_clean:
                    # Empty paragraph (extra spacing)
                    if para_idx < len(paragraphs_raw) - 1:  # Don't add trailing empty para
                        doc.add_paragraph()
                    continue
                
                # Create paragraph
                p = doc.add_paragraph()
                
                # Find segments that belong to this paragraph
                # Calculate character range for this paragraph in original text
                para_start_char = char_offset
                para_end_char = char_offset + len(para_text)
                
                # Extract segments within this paragraph range
                para_segments = []
                segment_char_pos = 0
                
                for segment in segments:
                    seg_text = segment['text']
                    seg_start = segment_char_pos
                    seg_end = segment_char_pos + len(seg_text)
                    
                    # Check if this segment overlaps with current paragraph
                    if seg_end > para_start_char and seg_start < para_end_char:
                        # Calculate overlap
                        overlap_start = max(0, para_start_char - seg_start)
                        overlap_end = min(len(seg_text), para_end_char - seg_start)
                        
                        if overlap_end > overlap_start:
                            # Extract the overlapping portion
                            overlap_text = seg_text[overlap_start:overlap_end]
                            # Clean newlines (replace with spaces for paragraph text)
                            overlap_text = overlap_text.replace('\n', ' ')
                            
                            if overlap_text.strip():
                                para_segments.append({
                                    'text': overlap_text,
                                    'tags': segment['tags']
                                })
                    
                    segment_char_pos += len(seg_text)
                
                # Build runs from segments, merging consecutive segments with same tags
                current_run_text = ""
                current_run_tags = None
                
                for seg in para_segments:
                    seg_text = seg['text']
                    seg_tags = seg['tags']
                    
                    # Check if tags changed
                    if current_run_tags != seg_tags and current_run_text:
                        # Finish current run
                        if current_run_text.strip():
                            run = p.add_run(current_run_text.strip())
                            apply_docx_formatting(run, current_run_tags, textbox, formatting_tags)
                        current_run_text = ""
                    
                    current_run_tags = seg_tags
                    current_run_text += seg_text
                
                # Add last run
                if current_run_text.strip():
                    run = p.add_run(current_run_text.strip())
                    apply_docx_formatting(run, current_run_tags, textbox, formatting_tags)
                
                # Update offset for next paragraph
                char_offset = para_end_char + 2  # +2 for the '\n\n' separator
                    
        except Exception as e:
            print(f"Error exporting with formatting: {e}, falling back to plain text")
            # Fall through to plain text export
    
    if not textbox_widget or not hasattr(textbox_widget, '_textbox'):
        # Simple export without formatting
        for paragraph in text.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
    elif formatting_tags:
        # Export with formatting
        # Split text into paragraphs
        paragraphs = text.split("\n\n")
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
                
            p = doc.add_paragraph()
            
            # Process text character by character with formatting
            char_index = 0
            current_run = None
            current_font_family = None
            current_font_size = None
            current_bold = False
            current_italic = False
            current_underline = False
            current_font_color = None
            current_highlight = None
            
            # Convert character positions to tkinter indices for tag lookup
            for i, char in enumerate(para_text):
                # Find all tags that apply to this character position
                # We need to map character position to textbox index
                # For simplicity, we'll check tags that might apply to this range
                char_tags = {}
                
                # Check formatting_tags for tags that might apply
                for tag_name, tag_info in formatting_tags.items():
                    try:
                        tag_type = tag_info.get("type")
                        start_str = tag_info.get("start", "")
                        end_str = tag_info.get("end", "")
                        
                        # Try to parse positions (could be "1.5" format or char offset)
                        try:
                            # Parse tkinter index format (e.g., "1.5" means line 1, char 5)
                            if "." in str(start_str):
                                parts = str(start_str).split(".")
                                start_line = int(parts[0]) - 1  # Convert to 0-based
                                start_char = int(parts[1]) if len(parts) > 1 else 0
                                # Approximate: we'll use paragraph start + char_index
                                # This is simplified - in real use, we'd need full text mapping
                                tag_start_char = start_char
                            else:
                                tag_start_char = int(start_str) if str(start_str).isdigit() else 0
                            
                            if "." in str(end_str):
                                parts = str(end_str).split(".")
                                end_line = int(parts[0]) - 1
                                end_char = int(parts[1]) if len(parts) > 1 else 0
                                tag_end_char = end_char
                            else:
                                tag_end_char = int(end_str) if str(end_str).isdigit() else len(para_text)
                            
                            # Check if current char position is within tag range
                            # Note: This is simplified - assumes tags are within current paragraph
                            if tag_start_char <= char_index < tag_end_char:
                                char_tags[tag_type] = tag_info
                        except:
                            pass
                    except:
                        pass
                
                # Determine formatting for this character
                font_family = None
                font_size = None
                bold = False
                italic = False
                underline = False
                font_color = None
                highlight = None
                
                # Merge all applicable tags
                for tag_type, tag_info in char_tags.items():
                    if tag_type == "font_merged":
                        font_family = tag_info.get("family", font_family)
                        font_size = tag_info.get("size", font_size)
                        bold = tag_info.get("bold", bold)
                        italic = tag_info.get("italic", italic)
                    elif tag_type == "bold":
                        bold = True
                    elif tag_type == "italic":
                        italic = True
                    elif tag_type == "underline":
                        underline = True
                    elif tag_type == "fontcolor":
                        font_color = tag_info.get("color", font_color)
                    elif tag_type == "highlight":
                        highlight = tag_info.get("color", highlight)
                    elif tag_type == "font":
                        font_family = tag_info.get("family", font_family)
                    elif tag_type == "fontsize":
                        font_size = tag_info.get("size", font_size)
                
                # Check if we need to start a new run (formatting changed)
                needs_new_run = (
                    font_family != current_font_family or
                    font_size != current_font_size or
                    bold != current_bold or
                    italic != current_italic or
                    underline != current_underline or
                    font_color != current_font_color or
                    highlight != current_highlight
                )
                
                if needs_new_run or current_run is None:
                    # Finish previous run
                    if current_run is not None:
                        pass  # Already added
                    
                    # Start new run
                    current_run = p.add_run(char)
                    current_font_family = font_family
                    current_font_size = font_size
                    current_bold = bold
                    current_italic = italic
                    current_underline = underline
                    current_font_color = font_color
                    current_highlight = highlight
                    
                    # Apply formatting
                    if current_bold:
                        current_run.bold = True
                    if current_italic:
                        current_run.italic = True
                    if current_underline:
                        current_run.underline = True
                    if current_font_size:
                        try:
                            current_run.font.size = Pt(int(current_font_size))
                        except:
                            pass
                    if current_font_family:
                        try:
                            current_run.font.name = current_font_family
                        except:
                            pass
                    if current_font_color:
                        try:
                            # Convert hex color to RGB
                            hex_color = current_font_color.lstrip('#')
                            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                            current_run.font.color.rgb = RGBColor(*rgb)
                        except:
                            pass
                    if current_highlight:
                        try:
                            # Convert hex color to RGB for highlight
                            hex_color = current_highlight.lstrip('#')
                            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                            current_run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # DOCX limited colors
                        except:
                            pass
                else:
                    # Continue current run - append to text
                    current_run.text += char
                
                char_index += 1
    
    doc.save(file_path)

def export_to_pdf(text, file_path, formatting_tags=None, textbox_widget=None):
    """Export text to PDF with formatting support."""
    from reportlab.platypus import Paragraph  # Import at function level to avoid scope issues
    
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Use textbox widget if available for better formatting extraction
    if textbox_widget and hasattr(textbox_widget, '_textbox'):
        try:
            textbox = textbox_widget._textbox
            segments = extract_formatted_text_segments(textbox)
            
            # Split full text into paragraphs by double newlines first
            full_text = textbox.get("1.0", "end-1c")
            paragraphs_raw = full_text.split('\n\n')
            
            # Process each paragraph
            char_offset = 0
            for para_idx, para_text in enumerate(paragraphs_raw):
                # Clean paragraph text (replace single newlines with spaces)
                para_text_clean = para_text.replace('\n', ' ').strip()
                
                if not para_text_clean:
                    # Empty paragraph (extra spacing)
                    if para_idx < len(paragraphs_raw) - 1:  # Don't add trailing empty para
                        story.append(Spacer(1, 12))
                    continue
                
                # Find segments that belong to this paragraph
                # Calculate character range for this paragraph in original text
                para_start_char = char_offset
                para_end_char = char_offset + len(para_text)
                
                # Extract segments within this paragraph range
                para_segments = []
                segment_char_pos = 0
                
                for segment in segments:
                    seg_text = segment['text']
                    seg_start = segment_char_pos
                    seg_end = segment_char_pos + len(seg_text)
                    
                    # Check if this segment overlaps with current paragraph
                    if seg_end > para_start_char and seg_start < para_end_char:
                        # Calculate overlap
                        overlap_start = max(0, para_start_char - seg_start)
                        overlap_end = min(len(seg_text), para_end_char - seg_start)
                        
                        if overlap_end > overlap_start:
                            # Extract the overlapping portion
                            overlap_text = seg_text[overlap_start:overlap_end]
                            # Clean newlines (replace with spaces for paragraph text)
                            overlap_text = overlap_text.replace('\n', ' ')
                            
                            if overlap_text.strip():
                                para_segments.append({
                                    'text': overlap_text,
                                    'tags': segment['tags']
                                })
                    
                    segment_char_pos += len(seg_text)
                
                # Build styled HTML from segments, merging consecutive segments with same tags
                para_html_parts = []
                current_html_text = ""
                current_html_tags = None
                
                for seg in para_segments:
                    seg_text = seg['text']
                    seg_tags = seg['tags']
                    
                    # Check if tags changed
                    if current_html_tags != seg_tags and current_html_text:
                        # Finish current styled segment
                        if current_html_text.strip():
                            styled = apply_pdf_formatting_styles(current_html_text.strip(), current_html_tags, textbox, formatting_tags)
                            para_html_parts.append(styled)
                        current_html_text = ""
                    
                    current_html_tags = seg_tags
                    current_html_text += seg_text
                
                # Add last segment
                if current_html_text.strip():
                    styled = apply_pdf_formatting_styles(current_html_text.strip(), current_html_tags, textbox, formatting_tags)
                    para_html_parts.append(styled)
                
                # Combine into paragraph HTML
                para_html = "".join(para_html_parts)
                if para_html.strip():
                    story.append(Paragraph(para_html, styles["Normal"]))
                    story.append(Spacer(1, 12))
                
                # Update offset for next paragraph
                char_offset = para_end_char + 2  # +2 for the '\n\n' separator
                    
        except Exception as e:
            print(f"Error exporting PDF with formatting: {e}, falling back to plain text")
            # Fall through to plain text export
    
    if not textbox_widget or not hasattr(textbox_widget, '_textbox'):
        # Simple export without formatting
        for paragraph in text.split("\n\n"):
            if paragraph.strip():
                story.append(Paragraph(paragraph.replace('\n', ' '), styles["Normal"]))
                story.append(Spacer(1, 12))
    
    doc.build(story)

def export_to_json(text, file_path, content_type="transcript"):
    data = {
        "type": content_type,
        "content": text,
        "paragraphs": [p.strip() for p in text.split("\n\n") if p.strip()]
    }
    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def export_content(textbox, content_type="transcript"):
    """Generic export function for any textbox content."""
    text = textbox.get("1.0", "end").strip()
    
    if not text:
        return
    
    file_path = filedialog.asksaveasfilename(
        title=f"Export {content_type.title()}",
        defaultextension=".txt",
        filetypes=[
            ("Text File", "*.txt"),
            ("Word Document", "*.docx"),
            ("PDF Document", "*.pdf"),
            ("JSON File", "*.json")
        ]
    )
    
    if not file_path:
        return
    
    try:
        if file_path.endswith(".txt"):
            export_to_txt(text, file_path)
        elif file_path.endswith(".docx"):
            export_to_docx(text, file_path)
        elif file_path.endswith(".pdf"):
            export_to_pdf(text, file_path)
        elif file_path.endswith(".json"):
            export_to_json(text, file_path, content_type)
    except Exception as e:
        print(f"Export error: {e}")

# --- Model Selection ---
def on_model_change(selected_model):
    """Reload Whisper model when user selects a different one."""
    global model, current_model_name
    
    if selected_model == current_model_name:
        return
    
    # Update status to show loading
    status_label.configure(text=f"Loading {selected_model} model...")
    app.update()
    
    try:
        model = WhisperModel(selected_model, device=DEVICE, compute_type=COMPUTE_TYPE)
        current_model_name = selected_model
        update_status_label()
        
        # Save model preference
        user_settings["whisper_model"] = selected_model
        save_settings(user_settings)
    except Exception as e:
        status_label.configure(text=f"Error loading model: {e}")

def update_status_label():
    """Update the status label with current device and model info."""
    if USE_CUDA:
        gpu_name = torch.cuda.get_device_name(0)
        status_text = f"🚀 GPU Mode - {gpu_name} | Model: {current_model_name}"
        status_color = "#4ADE80"  # Green
    else:
        status_text = f"⚡ CPU Mode (faster-whisper) | Model: {current_model_name}"
        status_color = "#3c847b"  # Custom color for CPU mode
    
    status_label.configure(text=status_text, text_color=status_color)

# --- Progress Bar Functions ---
def start_progress_indeterminate(label_text="Processing..."):
    """Start progress bar in indeterminate (animated) mode."""
    progress_frame.pack(side="left", padx=(15, 0))
    progress_label.configure(text=label_text)
    progress_bar.configure(mode="indeterminate")
    progress_bar.start()

def update_progress(value, label_text="Processing..."):
    """Update progress bar to a specific percentage (0-100). Thread-safe."""
    def _update():
        progress_frame.pack(side="left", padx=(15, 0))  # Ensure it's visible
        progress_bar.stop()
        progress_bar.configure(mode="determinate")
        progress_bar.set(value / 100)
        progress_label.configure(text=label_text)
    app.after(0, _update)

def finish_progress():
    """Set progress to 100% and then hide it."""
    def _finish():
        progress_bar.stop()
        progress_bar.configure(mode="determinate")
        progress_bar.set(1.0)
        progress_label.configure(text="Done!")
        # Hide after a short delay
        app.after(1500, hide_progress)
    app.after(0, _finish)

def hide_progress():
    """Hide the progress bar."""
    progress_bar.stop()
    progress_bar.set(0)
    progress_frame.pack_forget()

def set_ui_busy(busy):
    """Enable/disable UI elements during processing."""
    state = "disabled" if busy else "normal"
    model_dropdown.configure(state=state if USE_CUDA else "disabled")
    plus_button.configure(state=state)
    saved_transcripts_dropdown.configure(state=state)
    load_saved_btn.configure(state=state)
    delete_saved_btn.configure(state=state)
    saved_audio_dropdown.configure(state=state)
    load_audio_btn.configure(state=state)
    play_audio_btn.configure(state=state)
    delete_audio_btn.configure(state=state)
    # Disable/enable panel buttons
    for panel in panels:
        panel.export_btn.configure(state=state)
        panel.select_btn.configure(state=state)
        panel.minimize_btn.configure(state=state)
        panel.import_btn.configure(state=state)
        panel.save_btn.configure(state=state)
        panel.delete_btn.configure(state=state)

# --- Dynamic Panel System ---

# Available fonts for dropdown - 24 popular and useful fonts
AVAILABLE_FONTS = [
    # Sans-serif fonts (modern, clean)
    "Segoe UI", "Arial", "Verdana", "Tahoma", "Calibri", 
    "Helvetica", "Trebuchet MS", "Lucida Sans Unicode", "Century Gothic",
    "Franklin Gothic Medium", "Gill Sans MT", "Comic Sans MS",
    
    # Serif fonts (traditional, readable)
    "Times New Roman", "Georgia", "Garamond", "Book Antiqua", 
    "Palatino Linotype", "Baskerville Old Face", "Cambria", "Constantia",
    
    # Monospace fonts (code, technical)
    "Courier New", "Consolas", "Lucida Console", "Monaco"
]
AVAILABLE_FONT_SIZES = ["10", "11", "12", "13", "14", "16", "18", "20", "24", "28", "32"]

# Panel storage
panels = []  # List of TranscriptPanel objects
panel_counter = 0

class TranscriptPanel:
    """A dynamic transcript panel that can be minimized/expanded, reordered, and resized."""
    
    def __init__(self, parent_frame, panel_id, label_text="Panel"):
        self.panel_id = panel_id
        self.label_text = label_text
        self.is_minimized = False
        self.content = ""
        self.font_settings = {"family": FONT_FAMILY, "size": int(FONT_SIZES["body"])}
        self.parent_frame = parent_frame
        self.saved_width = 450  # Initialize saved width for restore
        self.drag_start_x = 0
        self.is_dragging = False
        self.panel_width = 450  # Default width
        self.resize_start_x = 0
        self.resize_start_width = 450
        self.is_resizing = False
        self.resize_indicator = None
        self.associated_saved_file = None  # Track which saved transcript file this panel is linked to
        self.auto_save_timer = None  # Timer for debounced auto-save
        self.formatting_tags = {}  # Store formatting tags (highlight, underline, font changes)
        self.search_dialog = None  # Search dialog reference
        
        # Main container frame (with minimum width)
        self.container = ctk.CTkFrame(parent_frame, fg_color="transparent")
        self.container.pack(side="left", fill="both", expand=True, padx=5)
        self.container.configure(width=self.panel_width)
        self.container.pack_propagate(False)  # Prevent auto-resize during drag
        self.normal_padx = 5  # Store normal padding
        
        # Full panel frame (shown when not minimized)
        self.full_frame = ctk.CTkFrame(self.container, fg_color="transparent")
        self.full_frame.pack(fill="both", expand=True)
        
        # Header with label and controls (top row)
        self.header = ctk.CTkFrame(self.full_frame, fg_color="transparent")
        self.header.pack(fill="x", pady=(0, 3))
        
        # Drag handle (left side of header) - visible move button for rearranging panels
        # Using a frame with label inside for better visibility and drag handling
        self.drag_handle_frame = ctk.CTkFrame(
            self.header,
            width=28,
            height=24,
            fg_color="#444444",
            corner_radius=4,
            cursor="fleur"
        )
        self.drag_handle_frame.pack(side="left", padx=(0, 5))
        self.drag_handle_frame.pack_propagate(False)
        
        self.drag_handle = ctk.CTkLabel(
            self.drag_handle_frame,
            text="⋮⋮",
            font=(FONT_FAMILY, 14),
            text_color="#CCCCCC"
        )
        self.drag_handle.pack(expand=True)
        
        # Bind drag events to both frame and label
        self.drag_handle_frame.bind("<Button-1>", self.on_drag_start)
        self.drag_handle_frame.bind("<B1-Motion>", self.on_drag_motion)
        self.drag_handle_frame.bind("<ButtonRelease-1>", self.on_drag_end)
        self.drag_handle.bind("<Button-1>", self.on_drag_start)
        self.drag_handle.bind("<B1-Motion>", self.on_drag_motion)
        self.drag_handle.bind("<ButtonRelease-1>", self.on_drag_end)
        
        # Add hover effect
        def on_enter(e):
            self.drag_handle_frame.configure(fg_color="#666666")
        def on_leave(e):
            self.drag_handle_frame.configure(fg_color="#444444")
        self.drag_handle_frame.bind("<Enter>", on_enter)
        self.drag_handle_frame.bind("<Leave>", on_leave)
        self.drag_handle.bind("<Enter>", on_enter)
        self.drag_handle.bind("<Leave>", on_leave)
        
        self.label = ctk.CTkLabel(
            self.header,
            text=label_text,
            font=(FONT_FAMILY, FONT_SIZES["heading"], "bold")
        )
        self.label.pack(side="left")
        
        # Delete button (X) - far right
        self.delete_btn = ctk.CTkButton(
            self.header,
            text="✕",
            width=28,
            height=24,
            font=(FONT_FAMILY, 12),
            corner_radius=4,
            fg_color="#CC4444",
            hover_color="#FF5555",
            text_color="white",
            command=self.delete_panel
        )
        self.delete_btn.pack(side="right", padx=(5, 0))
        
        # Minimize button
        self.minimize_btn = ctk.CTkButton(
            self.header,
            text="—",
            width=28,
            height=24,
            font=(FONT_FAMILY, 14),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.minimize
        )
        self.minimize_btn.pack(side="right", padx=(5, 0))
        
        # Single controls row: All buttons and formatting options in one row (auto-wraps when panel resized)
        self.controls_container = ctk.CTkFrame(self.full_frame, fg_color="transparent")
        self.controls_container.pack(fill="x", pady=(0, 5))
        
        # Main controls row - all items will flow left to right, wrapping naturally when needed
        self.controls_row = ctk.CTkFrame(self.controls_container, fg_color="transparent")
        self.controls_row.pack(fill="x")
        
        # Import Audio button (left side)
        self.import_btn = ctk.CTkButton(
            self.controls_row,
            text="📁 Import",
            width=80,
            height=26,
            font=(FONT_FAMILY, FONT_SIZES["small"]),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.import_audio
        )
        self.import_btn.pack(side="left", padx=(0, 5))
        
        # Font dropdown
        self.font_dropdown = ctk.CTkOptionMenu(
            self.controls_row,
            values=AVAILABLE_FONTS,
            command=self.on_font_change,
            width=120,
            height=26,
            font=(FONT_FAMILY, FONT_SIZES["small"]),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            button_color=BUTTON_COLOR,
            button_hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR
        )
        self.font_dropdown.set(self.font_settings["family"])
        self.font_dropdown.pack(side="left", padx=(0, 5))
        
        # Font size dropdown
        self.font_size_dropdown = ctk.CTkOptionMenu(
            self.controls_row,
            values=AVAILABLE_FONT_SIZES,
            command=self.on_font_size_change,
            width=70,
            height=26,
            font=(FONT_FAMILY, FONT_SIZES["small"]),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            button_color=BUTTON_COLOR,
            button_hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR
        )
        self.font_size_dropdown.set(str(self.font_settings["size"]))
        self.font_size_dropdown.pack(side="left", padx=(0, 5))
        
        # Highlight color dropdown
        highlight_label = ctk.CTkLabel(self.controls_row, text="Highlight:", font=(FONT_FAMILY, FONT_SIZES["small"]))
        highlight_label.pack(side="left", padx=(5, 5))
        
        highlight_colors = [
            ("None", None),
            ("Yellow", "#FFFF00"),
            ("Green", "#90EE90"),
            ("Blue", "#ADD8E6"),
            ("Orange", "#FFA500"),
            ("Plum", "#DDA0DD"),
            ("Pink", "#FFB6C1"),
        ]
        
        highlight_values = [name for name, _ in highlight_colors]
        self.highlight_colors_dict = {name: color for name, color in highlight_colors}
        
        def on_highlight_select(choice):
            """Handle highlight color selection."""
            color = self.highlight_colors_dict.get(choice)
            if color is None:
                self.clear_highlight()
            else:
                self.apply_highlight(color)
        
        highlight_dropdown = ctk.CTkOptionMenu(
            self.controls_row,
            values=highlight_values,
            command=on_highlight_select,
            width=120,
            height=26,
            font=(FONT_FAMILY, FONT_SIZES["small"]),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            button_color=BUTTON_COLOR,
            button_hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR
        )
        highlight_dropdown.set("None")
        highlight_dropdown.pack(side="left", padx=(0, 5))
        
        # Bold button
        bold_btn = ctk.CTkButton(
            self.controls_row,
            text="B",
            width=28,
            height=26,
            font=(FONT_FAMILY, 14, "bold"),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.toggle_bold
        )
        bold_btn.pack(side="left", padx=(0, 5))
        
        # Italic button
        italic_btn = ctk.CTkButton(
            self.controls_row,
            text="I",
            width=28,
            height=26,
            font=(FONT_FAMILY, 14, "italic"),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.toggle_italic
        )
        italic_btn.pack(side="left", padx=(0, 5))
        
        # Underline button
        underline_btn = ctk.CTkButton(
            self.controls_row,
            text="U",
            width=28,
            height=26,
            font=(FONT_FAMILY, 12, "underline"),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.toggle_underline
        )
        underline_btn.pack(side="left", padx=(0, 5))
        
        # Font color button
        font_color_btn = ctk.CTkButton(
            self.controls_row,
            text="A",
            width=28,
            height=26,
            font=(FONT_FAMILY, 14),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.change_font_color
        )
        font_color_btn.pack(side="left", padx=(0, 5))
        
        # Search button
        search_btn = ctk.CTkButton(
            self.controls_row,
            text="🔍",
            width=28,
            height=26,
            font=(FONT_FAMILY, 12),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.show_search_dialog
        )
        search_btn.pack(side="left", padx=(5, 5))
        
        # Export button (right side)
        self.export_btn = ctk.CTkButton(
            self.controls_row,
            text="Export",
            width=60,
            height=26,
            font=(FONT_FAMILY, FONT_SIZES["small"]),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.export_content
        )
        self.export_btn.pack(side="right", padx=(5, 0))
        
        # Select All button
        self.select_btn = ctk.CTkButton(
            self.controls_row,
            text="Select All",
            width=70,
            height=26,
            font=(FONT_FAMILY, FONT_SIZES["small"]),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.select_all
        )
        self.select_btn.pack(side="right", padx=(5, 0))
        
        # Save button (with disk icon)
        self.save_btn = ctk.CTkButton(
            self.controls_row,
            text="💾",
            width=32,
            height=26,
            font=(FONT_FAMILY, 14),
            corner_radius=4,
            fg_color=BUTTON_COLOR,
            hover_color=BUTTON_HOVER_COLOR,
            text_color=BUTTON_TEXT_COLOR,
            command=self.save_transcript
        )
        self.save_btn.pack(side="right", padx=(5, 0))
        
        # Store reference for minimize/restore functionality
        self.formatting_row = self.controls_row  # Keep for compatibility
        
        # Content frame (holds textbox + resize handle)
        self.content_frame = ctk.CTkFrame(self.full_frame, fg_color="transparent")
        self.content_frame.pack(fill="both", expand=True)
        
        # Textbox
        self.textbox = ctk.CTkTextbox(
            self.content_frame,
            width=450,
            height=700,
            wrap="word",
            font=(FONT_FAMILY, FONT_SIZES["body"]),
            corner_radius=10,
            border_width=1,
            fg_color=get_theme_color("textbox_bg"),
            border_color=get_theme_color("textbox_border"),
            text_color=get_theme_color("textbox_text")
        )
        self.textbox.pack(side="left", fill="both", expand=True)
        
        # Bind to text changes for auto-save
        self.textbox._textbox.bind("<KeyRelease>", self.on_text_change)
        self.textbox._textbox.bind("<ButtonRelease>", self.on_text_change)  # For paste operations
        
        # Resize handle (right edge)
        self.resize_handle = ctk.CTkFrame(
            self.content_frame,
            width=6,
            fg_color=get_theme_color("resize_handle"),
            corner_radius=3,
            cursor="sb_h_double_arrow"
        )
        self.resize_handle.pack(side="right", fill="y", padx=(2, 0))
        
        # Bind resize events
        self.resize_handle.bind("<Button-1>", self.on_resize_start)
        self.resize_handle.bind("<B1-Motion>", self.on_resize_motion)
        self.resize_handle.bind("<ButtonRelease-1>", self.on_resize_end)
        self.resize_handle.bind("<Enter>", lambda e: self.resize_handle.configure(fg_color=BUTTON_COLOR))
        self.resize_handle.bind("<Leave>", lambda e: self.resize_handle.configure(fg_color=get_theme_color("resize_handle")) if not self.is_resizing else None)
        
        # Bind to container resize to handle control wrapping
        self.container.bind("<Configure>", self.on_container_resize)
        
        # Drag indicator label (shown during drag)
        self.drag_indicator = None
        
        # Minimized bar (hidden initially)
        self.mini_bar = ctk.CTkFrame(
            self.container,
            width=30,
            height=700,
            fg_color=get_theme_color("mini_bar"),
            corner_radius=8
        )
        
        # Label on minimized bar (rotated text simulation)
        self.mini_label = ctk.CTkLabel(
            self.mini_bar,
            text=label_text[:15] + "..." if len(label_text) > 15 else label_text,
            font=(FONT_FAMILY, FONT_SIZES["small"]),
            text_color="#888888"
        )
        self.mini_label.place(relx=0.5, rely=0.5, anchor="center")
        
        # Preview frame (shown on hover)
        self.preview_frame = ctk.CTkFrame(
            self.container,
            width=200,
            height=700,
            fg_color=get_theme_color("textbox_bg"),
            corner_radius=8,
            border_width=1,
            border_color=BUTTON_COLOR
        )
        
        self.preview_label = ctk.CTkLabel(
            self.preview_frame,
            text="",
            font=(FONT_FAMILY, FONT_SIZES["small"]),
            wraplength=180,
            justify="left"
        )
        self.preview_label.pack(padx=10, pady=10, fill="both", expand=True)
        
        # Bind hover events to minimized bar
        self.mini_bar.bind("<Enter>", self.on_hover_enter)
        self.mini_bar.bind("<Leave>", self.on_hover_leave)
        self.mini_bar.bind("<Button-1>", self.on_click_restore)
        self.mini_label.bind("<Enter>", self.on_hover_enter)
        self.mini_label.bind("<Leave>", self.on_hover_leave)
        self.mini_label.bind("<Button-1>", self.on_click_restore)
    
    def minimize(self):
        """Minimize panel to a vertical bar."""
        self.is_minimized = True
        self.content = self.textbox.get("1.0", "end").strip()
        self.full_frame.pack_forget()
        self.formatting_row.pack_forget()  # Hide formatting toolbar when minimized
        self.mini_bar.pack(fill="y", expand=False, pady=5)
        
        # Store the current width before minimizing
        self.saved_width = self.panel_width
        
        # Set container to minimal width and reduce padding to stack close
        self.container.configure(width=40)
        self.container.pack_configure(padx=2)  # Reduce padding when minimized
        self.container.update_idletasks()  # Force update
    
    def restore(self):
        """Restore panel from minimized state."""
        self.is_minimized = False
        self.mini_bar.pack_forget()
        self.preview_frame.pack_forget()
        self.full_frame.pack(fill="both", expand=True)
        self.controls_container.pack(fill="x", pady=(0, 5))  # Show controls when restored
        
        # Restore the saved width (or default to 450)
        restore_width = getattr(self, 'saved_width', 450)
        self.panel_width = restore_width
        self.container.configure(width=restore_width)
        self.container.pack_configure(padx=self.normal_padx)  # Restore normal padding
        self.container.update_idletasks()  # Force update
    
    def on_hover_enter(self, event):
        """Show preview on hover."""
        if self.is_minimized:
            preview_text = self.content[:500] + "..." if len(self.content) > 500 else self.content
            self.preview_label.configure(text=preview_text if preview_text else "(Empty)")
            self.mini_bar.pack_forget()
            self.preview_frame.pack(fill="y", expand=False, pady=5)
            # Bind events to preview frame too
            self.preview_frame.bind("<Leave>", self.on_preview_leave)
            self.preview_frame.bind("<Button-1>", self.on_click_restore)
            self.preview_label.bind("<Button-1>", self.on_click_restore)
    
    def on_hover_leave(self, event):
        """Collapse preview on mouse leave from mini bar."""
        # Don't collapse if we're entering the preview frame
        pass
    
    def on_preview_leave(self, event):
        """Collapse preview when leaving preview frame."""
        if self.is_minimized:
            # Check if mouse actually left the preview area
            widget = event.widget
            x, y = event.x, event.y
            width = widget.winfo_width()
            height = widget.winfo_height()
            # Only collapse if mouse truly exited
            if x <= 0 or x >= width or y <= 0 or y >= height:
                self.preview_frame.pack_forget()
                self.mini_bar.pack(fill="y", expand=False, pady=5)
    
    def on_click_restore(self, event):
        """Restore panel on click."""
        self.restore()
    
    def on_drag_start(self, event):
        """Start dragging the panel."""
        self.drag_start_x = event.x_root
        self.is_dragging = True
        
        # Visual feedback on container
        self.container.configure(border_width=2, border_color=BUTTON_COLOR)
        
        # Create vertical divider line indicator (turquoise)
        self.drag_indicator = ctk.CTkToplevel(app)
        self.drag_indicator.overrideredirect(True)
        self.drag_indicator.attributes("-alpha", 0.9)
        self.drag_indicator.attributes("-topmost", True)
        
        # Get panels container position for indicator placement
        container_y = self.parent_frame.winfo_rooty()
        container_height = self.parent_frame.winfo_height()
        
        # Create the vertical divider line (turquoise)
        divider_line = ctk.CTkFrame(
            self.drag_indicator,
            width=4,
            height=container_height,
            fg_color=BUTTON_COLOR,  # Turquoise color
            corner_radius=2
        )
        divider_line.pack(fill="both", expand=True)
        
        # Position at current mouse x position (centered on cursor)
        self.drag_indicator.geometry(f"4x{container_height}+{event.x_root - 2}+{container_y}")
        
        # Dim other panels slightly
        for panel in panels:
            if panel != self:
                panel.container.configure(fg_color="#1a1a1a")
    
    def on_drag_motion(self, event):
        """Handle drag motion for reordering."""
        if self.is_dragging and self.drag_indicator:
            # Move the vertical divider line with cursor (centered on cursor)
            container_y = self.parent_frame.winfo_rooty()
            container_height = self.parent_frame.winfo_height()
            self.drag_indicator.geometry(f"4x{container_height}+{event.x_root - 2}+{container_y}")
            
            # Calculate potential new position and highlight drop zone
            delta = event.x_root - self.drag_start_x
            try:
                current_idx = panels.index(self)
                panel_width = self.panel_width
                positions_moved = int(delta / (panel_width / 2))
                new_idx = max(0, min(len(panels) - 1, current_idx + positions_moved))
                
                # Highlight potential drop position
                for i, panel in enumerate(panels):
                    if panel != self:
                        if i == new_idx:
                            panel.container.configure(fg_color="#2a3a3a", border_width=1, border_color=BUTTON_COLOR)
                        else:
                            panel.container.configure(fg_color="#1a1a1a", border_width=0)
            except ValueError:
                pass
    
    def on_drag_end(self, event):
        """End dragging and reorder panels if needed."""
        self.is_dragging = False
        self.container.configure(border_width=0)
        
        # Destroy floating indicator
        if self.drag_indicator:
            self.drag_indicator.destroy()
            self.drag_indicator = None
        
        # Reset all panel backgrounds
        for panel in panels:
            panel.container.configure(fg_color="transparent", border_width=0)
        
        # Calculate which position to move to based on x position
        current_x = event.x_root
        delta = current_x - self.drag_start_x
        
        # Get current index in panels list
        try:
            current_idx = panels.index(self)
        except ValueError:
            return
        
        # Determine new position based on drag distance
        panel_width = self.panel_width
        positions_moved = int(delta / (panel_width / 2))
        
        if positions_moved != 0:
            new_idx = max(0, min(len(panels) - 1, current_idx + positions_moved))
            if new_idx != current_idx:
                # Reorder in list
                panels.remove(self)
                panels.insert(new_idx, self)
                
                # Repack all panels in new order
                for panel in panels:
                    panel.container.pack_forget()
                for panel in panels:
                    # Use appropriate padding based on minimized state
                    padx = 2 if panel.is_minimized else panel.normal_padx
                    panel.container.pack(side="left", fill="both", expand=True, padx=padx)
                
                # Don't renumber panels - labels stay with their panels
    
    def on_container_resize(self, event):
        """Handle container resize - placeholder for future wrapping functionality."""
        # This is called when container is resized
        # Future: could implement control wrapping logic here
        pass
    
    def on_resize_start(self, event):
        """Start resizing the panel."""
        self.resize_start_x = event.x_root
        self.resize_start_width = self.panel_width
        self.is_resizing = True
        self.resize_handle.configure(fg_color=BUTTON_COLOR)
        
        # Create a vertical line indicator for resize preview
        self.resize_indicator = ctk.CTkToplevel(app)
        self.resize_indicator.overrideredirect(True)
        self.resize_indicator.attributes("-alpha", 0.8)
        self.resize_indicator.attributes("-topmost", True)
        
        # Get container position for indicator placement
        container_y = self.container.winfo_rooty()
        container_height = self.container.winfo_height()
        
        # Create the indicator line
        indicator_line = ctk.CTkFrame(
            self.resize_indicator,
            width=4,
            height=container_height,
            fg_color=BUTTON_COLOR,
            corner_radius=2
        )
        indicator_line.pack(fill="both", expand=True)
        
        # Position centered on cursor
        self.resize_indicator.geometry(f"4x{container_height}+{event.x_root - 2}+{container_y}")
    
    def on_resize_motion(self, event):
        """Handle resize motion - update preview indicator centered on cursor."""
        if self.is_resizing and self.resize_indicator:
            delta = event.x_root - self.resize_start_x
            new_width = max(200, min(900, self.resize_start_width + delta))
            
            # Move indicator centered on cursor position
            container_y = self.container.winfo_rooty()
            container_height = self.container.winfo_height()
            
            # Center the 4px wide indicator on the cursor (cursor_x - 2)
            self.resize_indicator.geometry(f"4x{container_height}+{event.x_root - 2}+{container_y}")
    
    def on_resize_end(self, event):
        """End resizing and apply final width."""
        if self.is_resizing:
            # Destroy preview indicator
            if self.resize_indicator:
                self.resize_indicator.destroy()
                self.resize_indicator = None
            
            # Calculate and apply final width
            delta = event.x_root - self.resize_start_x
            new_width = max(200, min(900, self.resize_start_width + delta))
            
            # Apply the new width in one update
            self.panel_width = new_width
            self.saved_width = new_width  # Save width for restore
            self.container.configure(width=new_width)
            self.textbox.configure(width=new_width - 20)  # Account for padding/handle
            
            self.is_resizing = False
            self.resize_handle.configure(fg_color=get_theme_color("resize_handle"))
            
            # Force a single redraw
            self.container.update_idletasks()
    
    def import_audio(self):
        """Import audio file and transcribe to this panel."""
        file_path = filedialog.askopenfilename(
            title="Select Audio File",
            filetypes=[("MP3 Files", "*.mp3"), ("WAV Files", "*.wav"), ("All Audio", "*.mp3;*.wav")]
        )
        if file_path:
            self.transcribe_to_panel(file_path)
    
    def on_font_change(self, selected_font):
        """Handle font family change."""
        # Check if text is selected - if so, change only selection, otherwise change all
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            if sel_start and sel_end and sel_start != sel_end:
                # Change selected text only
                self.change_selected_font(selected_font)
            else:
                # Change all text
                self.font_settings["family"] = selected_font
                self.update_textbox_font()
        except:
            # No selection - change all
            self.font_settings["family"] = selected_font
            self.update_textbox_font()
    
    def on_font_size_change(self, selected_size):
        """Handle font size change."""
        # Check if text is selected - if so, change only selection, otherwise change all
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            if sel_start and sel_end and sel_start != sel_end:
                # Change selected text only
                self.change_selected_font_size(selected_size)
            else:
                # Change all text
                self.font_settings["size"] = int(selected_size)
                self.update_textbox_font()
        except:
            # No selection - change all
            self.font_settings["size"] = int(selected_size)
            self.update_textbox_font()
    
    def update_textbox_font(self):
        """Update textbox font based on current font settings."""
        font_size = int(self.font_settings["size"])  # Ensure it's an integer
        new_font = (self.font_settings["family"], font_size)
        self.textbox.configure(font=new_font)
        
        # Refresh text if there's content
        if hasattr(self, 'content') and self.content:
            self.set_text(self.content)
    
    def transcribe_to_panel(self, file_path):
        """Transcribe audio file to this specific panel."""
        if not os.path.exists(file_path):
            self.set_text(f"Error: File not found.\n\n{file_path}")
            return
        
        # Capture self reference for use in worker thread
        panel_self = self
        
        # Start transcription in background
        def worker():
            try:
                # Show progress bar
                app.after(0, lambda: start_progress_indeterminate("Transcribing audio..."))
                panel_self.set_text("Transcribing audio... Please wait.")
                
                # faster-whisper transcription
                # Disable condition_on_previous_text to prevent cascading errors in long transcriptions
                # This ensures each segment is transcribed independently, preventing jumbling
                segments, info = model.transcribe(
                    file_path, 
                    beam_size=5,
                    condition_on_previous_text=False  # Prevents cascading errors in long transcriptions
                )
                
                # Update progress - processing segments
                app.after(0, lambda: update_progress(50, "Processing transcription..."))
                
                # Collect all segments with proper ordering
                text_parts = []
                for segment in segments:
                    if segment.text and segment.text.strip():
                        text_parts.append(segment.text.strip())
                
                # Join segments with spaces
                text = " ".join(text_parts)
                
                # Update progress - formatting
                app.after(0, lambda: update_progress(75, "Formatting transcript..."))
                
                # Format transcript into paragraphs
                formatted = format_transcript(text)
                
                # Update progress - highlighting
                # Update progress - complete
                app.after(0, lambda: finish_progress())
                
                # Update panel with formatted text
                app.after(0, lambda t=formatted: panel_self.set_text(t))
                
            except Exception as e:
                app.after(0, lambda: hide_progress())
                app.after(0, lambda err=str(e): panel_self.set_text(f"Error during transcription:\n\n{err}"))
        
        thread = threading.Thread(target=worker, daemon=True)
        thread.start()
    
    def save_transcript(self):
        """Save transcript to in-app folder. Preserves formatting tags."""
        # Get current text from textbox
        text = self.get_text()
        
        if not text.strip():
            return
        
        # Ask for filename
        from datetime import datetime
        default_name = f"transcript_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Simple dialog for filename
        dialog = ctk.CTkInputDialog(
            text="Enter a name for this transcript:",
            title="Save Transcript"
        )
        filename = dialog.get_input()
        
        if filename:
            # Save text with formatting tags
            saved_name = save_transcript_to_app(text, filename, self.formatting_tags)
            if saved_name:
                # Associate this panel with the saved file for auto-save
                self.associated_saved_file = saved_name
                refresh_saved_transcripts_dropdown()
    
    def delete_panel(self):
        """Delete this panel from the UI."""
        global panels
        if len(panels) <= 1:
            # Don't delete the last panel
            return
        
        # Remove from panels list
        if self in panels:
            panels.remove(self)
        
        # Destroy the container
        self.destroy()
        
        # Renumber remaining panels
        renumber_panels()
    
    def select_all(self):
        """Select all text in the textbox."""
        self.textbox.focus_set()
        self.textbox._textbox.tag_add("sel", "1.0", "end")
    
    def apply_highlight(self, color):
        """Apply highlight color to selected text - allows stacking with other formatting."""
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            
            if sel_start and sel_end and sel_start != sel_end:
                self.apply_formatting_tag("highlight", value=color, tag_config={"background": color})
        except Exception as e:
            print(f"Error applying highlight: {e}")
    
    def clear_highlight(self):
        """Clear highlight from selected text."""
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            
            if sel_start and sel_end:
                # Remove all highlight tags in selection
                for tag_name in list(self.formatting_tags.keys()):
                    if tag_name.startswith("highlight_"):
                        self.textbox._textbox.tag_remove(tag_name, sel_start, sel_end)
                        # Update stored tags
                        if tag_name in self.formatting_tags:
                            del self.formatting_tags[tag_name]
                self.on_text_change()  # Trigger save
        except Exception as e:
            print(f"Error clearing highlight: {e}")
    
    def get_current_font_properties(self, sel_start):
        """Get current font properties from existing tags at a position."""
        base_font = self.textbox.cget("font")
        if isinstance(base_font, tuple):
            font_family = base_font[0]
            font_size = int(base_font[1]) if len(base_font) > 1 else int(self.font_settings["size"])
            styles = list(base_font[2:]) if len(base_font) > 2 else []
        else:
            font_family = self.font_settings["family"]
            font_size = int(self.font_settings["size"])
            styles = []
        
        # Check existing tags for font properties
        existing_tags = self.textbox._textbox.tag_names(sel_start)
        for tag in existing_tags:
            try:
                tag_config = self.textbox._textbox.tag_config(tag)
                if 'font' in tag_config and tag_config['font'][4]:  # [4] is the actual value
                    existing_font = tag_config['font'][4]
                    if isinstance(existing_font, tuple):
                        if len(existing_font) > 0:
                            font_family = existing_font[0]
                        if len(existing_font) > 1:
                            font_size = int(existing_font[1])
                        if len(existing_font) > 2:
                            styles = list(existing_font[2:])
            except:
                pass
        
        return font_family, font_size, styles
    
    def merge_and_apply_font_tag(self, sel_start, sel_end, **font_props):
        """Merge font properties and apply a single font tag (since tkinter tags don't stack fonts)."""
        try:
            # Get current font properties
            font_family, font_size, styles = self.get_current_font_properties(sel_start)
            
            # Apply new properties (only override if explicitly provided)
            if 'family' in font_props and font_props['family']:
                font_family = font_props['family']
            if 'size' in font_props and font_props['size'] is not None:
                font_size = int(font_props['size'])
            if 'bold' in font_props:
                if font_props['bold']:
                    if 'bold' not in styles:
                        styles.append('bold')
                else:
                    if 'bold' in styles:
                        styles.remove('bold')
            if 'italic' in font_props:
                if font_props['italic']:
                    if 'italic' not in styles:
                        styles.append('italic')
                else:
                    if 'italic' in styles:
                        styles.remove('italic')
            
            # Create merged font tuple
            font_tuple = (font_family, font_size)
            if styles:
                font_tuple = tuple(list(font_tuple) + styles)
            
            # Remove existing font tags in this range
            existing_tags = self.textbox._textbox.tag_names(sel_start)
            for tag in existing_tags:
                try:
                    tag_config = self.textbox._textbox.tag_config(tag)
                    if 'font' in tag_config and tag_config['font'][4]:
                        # This is a font tag - remove it from selection
                        self.textbox._textbox.tag_remove(tag, sel_start, sel_end)
                        # Check if tag still exists elsewhere
                        ranges = self.textbox._textbox.tag_ranges(tag)
                        if not ranges or len(ranges) == 0:
                            if tag in self.formatting_tags:
                                del self.formatting_tags[tag]
                except:
                    pass
            
            # Create new merged font tag
            tag_name = f"font_merged_{len(self.formatting_tags)}_{id(sel_start)}"
            self.textbox._textbox.tag_configure(tag_name, font=font_tuple)
            self.textbox._textbox.tag_add(tag_name, sel_start, sel_end)
            
            # Store in formatting_tags
            tag_info = {
                "type": "font_merged",
                "start": str(sel_start),
                "end": str(sel_end),
                "family": font_family,
                "size": font_size,
                "bold": "bold" in styles,
                "italic": "italic" in styles
            }
            self.formatting_tags[tag_name] = tag_info
            
            return tag_name
        except Exception as e:
            print(f"Error merging font tag: {e}")
            return None
    
    def apply_formatting_tag(self, tag_type, value=None, tag_config=None):
        """Apply a formatting tag to selected text, allowing multiple tags to stack."""
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            
            if not sel_start or not sel_end or sel_start == sel_end:
                return
            
            # Non-font properties can stack (underline, color, highlight)
            if tag_type in ["underline", "fontcolor", "highlight"]:
                # Create a unique tag name
                tag_id = f"{tag_type}_{len(self.formatting_tags)}_{id(sel_start)}"
                tag_name = f"{tag_type}_{tag_id}"
                
                # Configure tag
                if tag_config:
                    self.textbox._textbox.tag_configure(tag_name, **tag_config)
                elif tag_type == "underline":
                    self.textbox._textbox.tag_configure(tag_name, underline=True)
                
                # Add tag to selection (will stack with existing tags)
                self.textbox._textbox.tag_add(tag_name, sel_start, sel_end)
                
                # Store in formatting_tags
                tag_info = {"type": tag_type, "start": str(sel_start), "end": str(sel_end)}
                if value is not None:
                    if tag_type == "fontcolor":
                        tag_info["color"] = value
                    elif tag_type == "highlight":
                        tag_info["color"] = value
                self.formatting_tags[tag_name] = tag_info
                
            # Font properties need to be merged into a single tag
            elif tag_type in ["bold", "italic", "fontsize", "font"]:
                # Get current font properties to preserve what's not being changed
                font_family, current_size, current_styles = self.get_current_font_properties(sel_start)
                
                font_props = {}
                if tag_type == "bold":
                    font_props['bold'] = True
                    # Preserve italic, size, and family
                    font_props['italic'] = "italic" in current_styles
                    font_props['size'] = current_size
                    font_props['family'] = font_family
                elif tag_type == "italic":
                    font_props['italic'] = True
                    # Preserve bold, size, and family
                    font_props['bold'] = "bold" in current_styles
                    font_props['size'] = current_size
                    font_props['family'] = font_family
                elif tag_type == "fontsize":
                    font_props['size'] = int(value) if value else current_size
                    # Preserve bold, italic, and family
                    font_props['bold'] = "bold" in current_styles
                    font_props['italic'] = "italic" in current_styles
                    font_props['family'] = font_family
                elif tag_type == "font":
                    font_props['family'] = value
                    # Preserve bold, italic, and size
                    font_props['bold'] = "bold" in current_styles
                    font_props['italic'] = "italic" in current_styles
                    font_props['size'] = current_size
                
                self.merge_and_apply_font_tag(sel_start, sel_end, **font_props)
            
            self.on_text_change()  # Trigger save
            
        except Exception as e:
            print(f"Error applying formatting tag {tag_type}: {e}")
    
    def toggle_bold(self):
        """Toggle bold on selected text - merges with other font properties."""
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            
            if sel_start and sel_end and sel_start != sel_end:
                # Check if text is already bold by examining font tags
                tags_at_pos = self.textbox._textbox.tag_names(sel_start)
                has_bold = False
                
                # Check font tags for bold
                for tag in tags_at_pos:
                    try:
                        tag_config = self.textbox._textbox.tag_config(tag)
                        if 'font' in tag_config and tag_config['font'][4]:
                            font_tuple = tag_config['font'][4]
                            if isinstance(font_tuple, tuple) and len(font_tuple) > 2:
                                has_bold = "bold" in font_tuple[2:]
                    except:
                        # Also check by tag name
                        if "bold" in str(tag) or "font_merged" in str(tag):
                            # Check formatting_tags
                            if tag in self.formatting_tags:
                                tag_info = self.formatting_tags[tag]
                                if tag_info.get("bold") or tag_info.get("type") == "bold":
                                    has_bold = True
                
                # Toggle bold
                if has_bold:
                    # Remove bold - merge font without bold
                    self.merge_and_apply_font_tag(sel_start, sel_end, bold=False)
                else:
                    # Add bold - merge font with bold
                    self.merge_and_apply_font_tag(sel_start, sel_end, bold=True)
                
                self.on_text_change()  # Trigger save
        except Exception as e:
            print(f"Error toggling bold: {e}")
    
    def toggle_italic(self):
        """Toggle italic on selected text - merges with other font properties."""
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            
            if sel_start and sel_end and sel_start != sel_end:
                # Check if text is already italic by examining font tags
                tags_at_pos = self.textbox._textbox.tag_names(sel_start)
                has_italic = False
                
                # Check font tags for italic
                for tag in tags_at_pos:
                    try:
                        tag_config = self.textbox._textbox.tag_config(tag)
                        if 'font' in tag_config and tag_config['font'][4]:
                            font_tuple = tag_config['font'][4]
                            if isinstance(font_tuple, tuple) and len(font_tuple) > 2:
                                has_italic = "italic" in font_tuple[2:]
                    except:
                        # Also check by tag name
                        if "italic" in str(tag) or "font_merged" in str(tag):
                            if tag in self.formatting_tags:
                                tag_info = self.formatting_tags[tag]
                                if tag_info.get("italic") or tag_info.get("type") == "italic":
                                    has_italic = True
                
                # Toggle italic
                if has_italic:
                    # Remove italic - merge font without italic
                    self.merge_and_apply_font_tag(sel_start, sel_end, italic=False)
                else:
                    # Add italic - merge font with italic
                    self.merge_and_apply_font_tag(sel_start, sel_end, italic=True)
                
                self.on_text_change()  # Trigger save
        except Exception as e:
            print(f"Error toggling italic: {e}")
    
    def toggle_underline(self):
        """Toggle underline on selected text - allows stacking with other formatting."""
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            
            if sel_start and sel_end and sel_start != sel_end:
                tags_at_pos = self.textbox._textbox.tag_names(sel_start)
                has_underline = any("underline" in str(tag) for tag in tags_at_pos)
                
                if has_underline:
                    # Remove underline from selection
                    for tag in tags_at_pos:
                        if "underline" in str(tag):
                            try:
                                self.textbox._textbox.tag_remove(tag, sel_start, sel_end)
                                ranges = self.textbox._textbox.tag_ranges(tag)
                                if not ranges or len(ranges) == 0:
                                    if tag in self.formatting_tags:
                                        del self.formatting_tags[tag]
                            except:
                                pass
                else:
                    # Add underline (can stack with other formatting)
                    self.apply_formatting_tag("underline", tag_config={"underline": True})
                
                self.on_text_change()  # Trigger save
        except Exception as e:
            print(f"Error toggling underline: {e}")
    
    def change_font_color(self):
        """Change font color of selected text - allows stacking with other formatting."""
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            
            if sel_start and sel_end and sel_start != sel_end:
                import tkinter.colorchooser as colorchooser
                color = colorchooser.askcolor(title="Choose Font Color")
                
                if color and color[1]:  # color[1] is hex string
                    self.apply_formatting_tag("fontcolor", value=color[1], tag_config={"foreground": color[1]})
        except Exception as e:
            print(f"Error changing font color: {e}")
    
    def change_selected_font(self, font_family):
        """Change font family of selected text - preserves other formatting."""
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            
            if sel_start and sel_end and sel_start != sel_end:
                # Change font family - merges with existing font properties
                self.apply_formatting_tag("font", value=font_family)
        except Exception as e:
            print(f"Error changing font: {e}")
    
    def change_selected_font_size(self, font_size):
        """Change font size of selected text - preserves other formatting."""
        try:
            sel_start = self.textbox._textbox.index("sel.first")
            sel_end = self.textbox._textbox.index("sel.last")
            
            if sel_start and sel_end and sel_start != sel_end:
                # Change font size - merges with existing font properties (bold, italic, etc.)
                # Get current font properties to preserve bold/italic
                font_family, current_size, styles = self.get_current_font_properties(sel_start)
                
                # Merge font size change with existing properties
                self.merge_and_apply_font_tag(sel_start, sel_end, size=int(font_size), 
                                            bold="bold" in styles, italic="italic" in styles,
                                            family=font_family)
                self.on_text_change()  # Trigger save
        except Exception as e:
            print(f"Error changing font size: {e}")
    
    def find_all_occurrences(self, search_term):
        """Find all occurrences of search term in textbox."""
        occurrences = []
        if not search_term:
            return occurrences
        
        try:
            start_pos = "1.0"
            while True:
                pos = self.textbox._textbox.search(search_term, start_pos, "end", nocase=True)
                if not pos:
                    break
                
                end_pos = f"{pos} + {len(search_term)} chars"
                occurrences.append((pos, end_pos))
                
                # Move start position past this match
                start_pos = end_pos
        except Exception as e:
            print(f"Error finding occurrences: {e}")
        
        return occurrences
    
    def show_search_dialog(self):
        """Show search dialog for finding text with navigation."""
        try:
            # Close existing dialog if open
            if hasattr(self, 'search_dialog') and self.search_dialog:
                try:
                    self.search_dialog.destroy()
                except:
                    pass
            
            # Create custom search dialog window
            search_dialog = ctk.CTkToplevel(app)
            search_dialog.title("Find")
            search_dialog.geometry("450x180")  # Bigger dialog
            search_dialog.resizable(False, False)
            
            # Make dialog modal
            search_dialog.transient(app)
            search_dialog.grab_set()
            
            # Store reference
            self.search_dialog = search_dialog
            self.search_occurrences = []
            self.current_search_index = -1
            self.current_search_term = ""
            
            # Main frame
            main_frame = ctk.CTkFrame(search_dialog)
            main_frame.pack(fill="both", expand=True, padx=15, pady=15)
            
            # Search term entry
            search_label = ctk.CTkLabel(main_frame, text="Search for:", font=(FONT_FAMILY, FONT_SIZES["body"]))
            search_label.pack(anchor="w", pady=(10, 5))
            
            search_entry = ctk.CTkEntry(main_frame, width=400, height=32, font=(FONT_FAMILY, FONT_SIZES["body"]))
            search_entry.pack(fill="x", padx=10, pady=(0, 10))
            search_entry.focus()
            
            # Status label
            status_label = ctk.CTkLabel(main_frame, text="", font=(FONT_FAMILY, FONT_SIZES["small"]))
            status_label.pack(anchor="w", padx=10, pady=(0, 10))
            
            # Buttons frame
            buttons_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
            buttons_frame.pack(fill="x", padx=10, pady=(0, 10))
            
            def navigate_to_match(direction):
                """Navigate to next or previous match."""
                if not self.search_occurrences:
                    return
                
                if direction == 1:  # Next
                    self.current_search_index = (self.current_search_index + 1) % len(self.search_occurrences)
                elif direction == -1:  # Previous
                    self.current_search_index = (self.current_search_index - 1) % len(self.search_occurrences)
                
                # Highlight current match
                self.textbox._textbox.tag_remove("search_current", "1.0", "end")
                if self.current_search_index >= 0 and self.current_search_index < len(self.search_occurrences):
                    pos, end_pos = self.search_occurrences[self.current_search_index]
                    self.textbox._textbox.tag_add("search_current", pos, end_pos)
                    
                    # Scroll to current match
                    self.textbox._textbox.see(pos)
                    self.textbox._textbox.mark_set("insert", pos)
                    self.textbox._textbox.tag_add("sel", pos, end_pos)
                    
                    # Update status
                    status_label.configure(text=f"Match {self.current_search_index + 1} of {len(self.search_occurrences)}")
            
            def do_search():
                """Perform search and find all occurrences."""
                search_term = search_entry.get().strip()
                if not search_term:
                    status_label.configure(text="Please enter a search term.")
                    return
                
                # Find all occurrences
                self.search_occurrences = self.find_all_occurrences(search_term)
                self.current_search_term = search_term
                
                if not self.search_occurrences:
                    status_label.configure(text=f"'{search_term}' not found")
                    # Remove any previous highlights
                    self.textbox._textbox.tag_remove("search", "1.0", "end")
                    self.textbox._textbox.tag_remove("search_current", "1.0", "end")
                    self.current_search_index = -1
                else:
                    # Highlight all occurrences
                    self.textbox._textbox.tag_configure("search", background="#FFD700")  # Gold for all
                    self.textbox._textbox.tag_configure("search_current", background="#FFA500")  # Orange for current
                    self.textbox._textbox.tag_remove("search", "1.0", "end")
                    self.textbox._textbox.tag_remove("search_current", "1.0", "end")
                    
                    # Highlight all occurrences
                    for pos, end_pos in self.search_occurrences:
                        self.textbox._textbox.tag_add("search", pos, end_pos)
                    
                    # Navigate to first occurrence
                    self.current_search_index = 0
                    navigate_to_match(0)
            
            # Search button
            search_btn = ctk.CTkButton(
                buttons_frame,
                text="Find All",
                command=do_search,
                width=100,
                height=32,
                font=(FONT_FAMILY, FONT_SIZES["body"]),
                fg_color=BUTTON_COLOR,
                hover_color=BUTTON_HOVER_COLOR,
                text_color=BUTTON_TEXT_COLOR
            )
            search_btn.pack(side="left", padx=(0, 5))
            
            # Previous button
            prev_btn = ctk.CTkButton(
                buttons_frame,
                text="◄ Previous",
                command=lambda: navigate_to_match(-1),
                width=100,
                height=32,
                font=(FONT_FAMILY, FONT_SIZES["body"]),
                fg_color=BUTTON_COLOR,
                hover_color=BUTTON_HOVER_COLOR,
                text_color=BUTTON_TEXT_COLOR
            )
            prev_btn.pack(side="left", padx=(0, 5))
            
            # Next button
            next_btn = ctk.CTkButton(
                buttons_frame,
                text="Next ►",
                command=lambda: navigate_to_match(1),
                width=100,
                height=32,
                font=(FONT_FAMILY, FONT_SIZES["body"]),
                fg_color=BUTTON_COLOR,
                hover_color=BUTTON_HOVER_COLOR,
                text_color=BUTTON_TEXT_COLOR
            )
            next_btn.pack(side="left", padx=(0, 5))
            
            # Close button
            close_btn = ctk.CTkButton(
                buttons_frame,
                text="Close",
                command=search_dialog.destroy,
                width=80,
                height=32,
                font=(FONT_FAMILY, FONT_SIZES["body"]),
                fg_color=BUTTON_COLOR,
                hover_color=BUTTON_HOVER_COLOR,
                text_color=BUTTON_TEXT_COLOR
            )
            close_btn.pack(side="right")
            
            # Bind Enter key to search
            def on_enter_key(event):
                do_search()
            search_entry.bind("<Return>", on_enter_key)
            
            # Bind Escape to close
            def on_escape_key(event):
                search_dialog.destroy()
            search_dialog.bind("<Escape>", on_escape_key)
            
            # Clean up when dialog closes
            def on_close():
                self.search_dialog = None
                # Remove search highlights when dialog closes
                if hasattr(self, 'textbox') and self.textbox:
                    try:
                        self.textbox._textbox.tag_remove("search", "1.0", "end")
                        self.textbox._textbox.tag_remove("search_current", "1.0", "end")
                    except:
                        pass
            
            search_dialog.protocol("WM_DELETE_WINDOW", lambda: (on_close(), search_dialog.destroy()))
            
        except Exception as e:
            print(f"Error showing search dialog: {e}")
            import traceback
            traceback.print_exc()
    
    def search_text(self, search_term, start_pos="1.0"):
        """Search for text in the textbox (legacy method, now redirects to dialog)."""
        # This method is kept for backward compatibility but now opens the dialog
        self.show_search_dialog()
    
    def export_content(self):
        """Export panel content with formatting."""
        text = self.textbox.get("1.0", "end").strip()
        if not text:
            return
        
        # Update formatting tags with current positions before export
        self.update_formatting_tag_positions()
        
        file_path = filedialog.asksaveasfilename(
            title=f"Export {self.label_text}",
            defaultextension=".txt",
            filetypes=[
                ("Text File", "*.txt"),
                ("Word Document", "*.docx"),
                ("PDF Document", "*.pdf"),
                ("JSON File", "*.json")
            ]
        )
        
        if file_path:
            try:
                if file_path.endswith(".txt"):
                    export_to_txt(text, file_path)
                elif file_path.endswith(".docx"):
                    export_to_docx(text, file_path, self.formatting_tags, self.textbox)
                elif file_path.endswith(".pdf"):
                    export_to_pdf(text, file_path, self.formatting_tags, self.textbox)
                elif file_path.endswith(".json"):
                    export_to_json(text, file_path, self.label_text)
            except Exception as e:
                print(f"Export error: {e}")
                import tkinter.messagebox as messagebox
                messagebox.showerror("Export Error", f"Failed to export: {e}")
    
    def set_text(self, text, apply_bold=False, formatting_tags=None):
        """Set text content.
        
        Args:
            text: Text content to set
            apply_bold: Deprecated - kept for compatibility but no longer used
            formatting_tags: Optional dict of formatting tags to restore
        """
        # Temporarily disable auto-save during set_text to avoid saving while loading
        was_auto_saving = self.associated_saved_file is not None
        temp_file = self.associated_saved_file
        if was_auto_saving:
            self.associated_saved_file = None
        
        self.textbox.delete("1.0", "end")
        
        # Clear existing formatting tags
        self.formatting_tags = {}
        
        # Remove ** markers if present (treat all text the same way)
        if "**" in text:
            # Remove ** markers - treat all text uniformly
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        
        # Insert text normally
        self.textbox.insert("1.0", text)
        
        # Always store the original text with ** markers for saving
        self.content = text
        
        # Restore formatting tags if provided (after text is inserted)
        if formatting_tags:
            # Need to restore after text is fully inserted
            app.after(10, lambda: self.restore_formatting_tags(formatting_tags))
        
        # Re-enable auto-save if it was enabled
        if was_auto_saving:
            self.associated_saved_file = temp_file
    
    def get_font_for_tag(self, base_font, tag_info):
        """Get the appropriate font tuple for a formatting tag, preserving other formatting."""
        if isinstance(base_font, tuple):
            font_family = base_font[0]
            font_size = int(base_font[1]) if len(base_font) > 1 else int(self.font_settings["size"])
            font_styles = []
            if len(base_font) > 2:
                # Preserve existing styles
                for style in base_font[2:]:
                    if style not in ["bold", "italic"]:
                        font_styles.append(style)
        else:
            font_family = self.font_settings["family"]
            font_size = int(self.font_settings["size"])
            font_styles = []
        
        # Apply tag-specific formatting
        if tag_info.get("type") == "bold":
            if "bold" not in font_styles:
                font_styles.append("bold")
        elif tag_info.get("type") == "italic":
            if "italic" not in font_styles:
                font_styles.append("italic")
        elif tag_info.get("type") == "fontsize":
            font_size = int(tag_info.get("size", font_size))
        elif tag_info.get("type") == "font":
            font_family = tag_info.get("family", font_family)
        
        if font_styles:
            return tuple([font_family, font_size] + font_styles)
        else:
            return (font_family, font_size)
    
    def restore_formatting_tags(self, formatting_tags):
        """Restore formatting tags from saved metadata."""
        try:
            if not formatting_tags:
                return
            
            base_font = self.textbox.cget("font")
            
            # Build a map of positions to tags
            position_tags = {}  # Maps (start, end) tuples to list of tags
            
            for tag_name, tag_info in formatting_tags.items():
                try:
                    start = tag_info.get("start")
                    end = tag_info.get("end")
                    
                    if not start or not end:
                        continue
                    
                    # Convert to tkinter index format if needed
                    try:
                        start_index = str(start)
                        end_index = str(end)
                        # Validate positions
                        self.textbox._textbox.index(start_index)
                        self.textbox._textbox.index(end_index)
                    except:
                        # Try converting from character offset
                        try:
                            if isinstance(start, int):
                                start_index = self.textbox._textbox.index(f"1.0 + {start} chars")
                                end_index = self.textbox._textbox.index(f"1.0 + {end} chars")
                            else:
                                continue
                        except:
                            continue
                    
                    # Group tags by position range
                    pos_key = (start_index, end_index)
                    if pos_key not in position_tags:
                        position_tags[pos_key] = []
                    position_tags[pos_key].append((tag_name, tag_info))
                    
                except Exception as e:
                    print(f"Error processing tag {tag_name}: {e}")
                    continue
            
            # Apply tags grouped by position
            for (start_index, end_index), tags_list in position_tags.items():
                try:
                    # Collect all formatting properties for this range
                    font_family = None
                    font_size = None
                    is_bold = False
                    is_italic = False
                    is_underline = False
                    font_color = None
                    highlight_color = None
                    
                    for tag_name, tag_info in tags_list:
                        tag_type = tag_info.get("type")
                        
                        if tag_type == "font_merged":
                            # Merged font tag contains all font properties
                            font_family = tag_info.get("family", font_family)
                            size_val = tag_info.get("size", font_size)
                            if size_val:
                                font_size = int(size_val) if not isinstance(size_val, int) else size_val
                            is_bold = tag_info.get("bold", is_bold)
                            is_italic = tag_info.get("italic", is_italic)
                        elif tag_type == "bold":
                            is_bold = True
                        elif tag_type == "italic":
                            is_italic = True
                        elif tag_type == "underline":
                            is_underline = True
                        elif tag_type == "fontcolor":
                            font_color = tag_info.get("color")
                        elif tag_type == "highlight":
                            highlight_color = tag_info.get("color")
                        elif tag_type == "font":
                            font_family = tag_info.get("family", font_family)
                        elif tag_type == "fontsize":
                            size_val = tag_info.get("size")
                            if size_val:
                                font_size = int(size_val) if not isinstance(size_val, int) else size_val
                    
                    # Get base font properties for defaults
                    if isinstance(base_font, tuple):
                        default_family = base_font[0]
                        default_size = int(base_font[1]) if len(base_font) > 1 else int(self.font_settings["size"])
                    else:
                        default_family = self.font_settings["family"]
                        default_size = int(self.font_settings["size"])
                    
                    if font_family is None:
                        font_family = default_family
                    if font_size is None:
                        font_size = default_size
                    
                    # Build font tuple with styles
                    font_styles = []
                    if is_bold:
                        font_styles.append("bold")
                    if is_italic:
                        font_styles.append("italic")
                    
                    font_tuple = (font_family, font_size)
                    if font_styles:
                        font_tuple = tuple(list(font_tuple) + font_styles)
                    
                    # Apply formatting as separate tags (font merged, others stack)
                    # 1. Font properties (as merged tag)
                    if is_bold or is_italic or font_family != default_family or font_size != default_size:
                        tag_name_font = f"restored_font_{len(self.formatting_tags)}"
                        self.textbox._textbox.tag_configure(tag_name_font, font=font_tuple)
                        self.textbox._textbox.tag_add(tag_name_font, start_index, end_index)
                        self.formatting_tags[tag_name_font] = {
                            "type": "font_merged",
                            "start": start_index,
                            "end": end_index,
                            "family": font_family,
                            "size": font_size,
                            "bold": is_bold,
                            "italic": is_italic
                        }
                    
                    # 2. Underline (separate tag, can stack)
                    if is_underline:
                        tag_name_underline = f"restored_underline_{len(self.formatting_tags)}"
                        self.textbox._textbox.tag_configure(tag_name_underline, underline=True)
                        self.textbox._textbox.tag_add(tag_name_underline, start_index, end_index)
                        self.formatting_tags[tag_name_underline] = {
                            "type": "underline",
                            "start": start_index,
                            "end": end_index
                        }
                    
                    # 3. Font color (separate tag, can stack)
                    if font_color:
                        tag_name_color = f"restored_fontcolor_{len(self.formatting_tags)}"
                        self.textbox._textbox.tag_configure(tag_name_color, foreground=font_color)
                        self.textbox._textbox.tag_add(tag_name_color, start_index, end_index)
                        self.formatting_tags[tag_name_color] = {
                            "type": "fontcolor",
                            "start": start_index,
                            "end": end_index,
                            "color": font_color
                        }
                    
                    # 4. Highlight (separate tag, can stack)
                    if highlight_color:
                        tag_name_highlight = f"restored_highlight_{len(self.formatting_tags)}"
                        self.textbox._textbox.tag_configure(tag_name_highlight, background=highlight_color)
                        self.textbox._textbox.tag_add(tag_name_highlight, start_index, end_index)
                        self.formatting_tags[tag_name_highlight] = {
                            "type": "highlight",
                            "start": start_index,
                            "end": end_index,
                            "color": highlight_color
                        }
                        
                except Exception as e:
                    print(f"Error restoring tags at {start_index}-{end_index}: {e}")
                    continue
                    
        except Exception as e:
            print(f"Error restoring formatting tags: {e}")
    
    def get_text(self):
        """Get text content."""
        return self.textbox.get("1.0", "end").strip()
    
    def on_text_change(self, event=None):
        """Handle text changes - trigger auto-save with debouncing."""
        # Update content to match current text
        self.content = self.get_text()
        
        # Only auto-save if there's an associated saved file
        if self.associated_saved_file:
            # Cancel previous timer
            if self.auto_save_timer:
                app.after_cancel(self.auto_save_timer)
            # Schedule auto-save after 1 second of no changes
            self.auto_save_timer = app.after(1000, self.auto_save_transcript)
    
    def auto_save_transcript(self):
        """Auto-save transcript to associated saved file."""
        if not self.associated_saved_file:
            return
        
        try:
            # Always get the current text from the textbox to ensure edits are saved
            text = self.get_text()
            
            # Update formatting tags with current positions
            self.update_formatting_tag_positions()
            
            # Save text and formatting tags
            save_transcript_to_app(text, self.associated_saved_file, self.formatting_tags)
            
            # Update content to match saved text
            self.content = text
            
        except Exception as e:
            print(f"Error auto-saving transcript: {e}")
    
    def update_formatting_tag_positions(self):
        """Update formatting tag positions to reflect current text state."""
        try:
            # Get all tags currently in the textbox
            for tag_name in list(self.formatting_tags.keys()):
                try:
                    # Get current ranges for this tag
                    ranges = self.textbox._textbox.tag_ranges(tag_name)
                    if ranges and len(ranges) >= 2:
                        # Update with current start and end positions
                        self.formatting_tags[tag_name]["start"] = str(ranges[0])
                        self.formatting_tags[tag_name]["end"] = str(ranges[1])
                    elif tag_name in self.formatting_tags:
                        # Tag no longer exists, remove it
                        del self.formatting_tags[tag_name]
                except:
                    # Tag might not exist anymore
                    if tag_name in self.formatting_tags:
                        del self.formatting_tags[tag_name]
        except Exception as e:
            print(f"Error updating formatting tag positions: {e}")
    
    def destroy(self):
        """Remove the panel."""
        # Cancel any pending auto-save
        if self.auto_save_timer:
            app.after_cancel(self.auto_save_timer)
        self.container.destroy()

def renumber_panels():
    """Renumber all panels as Transcript 1, Transcript 2, etc. in their current order."""
    for idx, panel in enumerate(panels, start=1):
        new_label = f"Transcript {idx}"
        panel.label_text = new_label
        panel.label.configure(text=new_label)
        # Update minimized label too
        panel.mini_label.configure(text=new_label[:15] + "..." if len(new_label) > 15 else new_label)

def create_panel(label_text=None, insert_at_start=False):
    """Create a new transcript panel with automatic Transcript X naming."""
    global panel_counter
    panel_counter += 1
    
    # Create panel with temporary label (will be renumbered)
    # Always use Transcript X naming - ignore any provided label_text
    label_text = "Transcript 1"  # Temporary, will be renumbered
    
    panel = TranscriptPanel(panels_container, panel_counter, label_text)
    
    # Insert at beginning (leftmost position) or append to end
    if insert_at_start:
        panels.insert(0, panel)
        # Repack all panels to maintain order
        for p in panels:
            p.container.pack_forget()
        for p in panels:
            padx = 2 if p.is_minimized else p.normal_padx
            p.container.pack(side="left", fill="both", expand=True, padx=padx)
    else:
        panels.append(panel)
    
    # Renumber all panels to ensure correct order (Transcript 1, 2, 3, etc.)
    renumber_panels()
    
    return panel

def add_new_panel():
    """Add a new panel via the plus button. Inserts at first position (leftmost)."""
    create_panel(insert_at_start=True)

# --- Main Layout ---

# Row 0: Header with Logo and Status
header_frame = ctk.CTkFrame(app, fg_color="transparent")
header_frame.pack(fill="x", padx=15, pady=(15, 8))

# Load and display logo (maintain aspect ratio)
LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png")
if os.path.exists(LOGO_PATH):
    original_logo = Image.open(LOGO_PATH)
    orig_width, orig_height = original_logo.size
    target_height = 60
    aspect_ratio = orig_width / orig_height
    target_width = int(target_height * aspect_ratio)
    
    logo_image = ctk.CTkImage(
        light_image=original_logo,
        dark_image=original_logo,
        size=(target_width, target_height)
    )
    logo_label = ctk.CTkLabel(header_frame, image=logo_image, text="")
    logo_label.pack(side="left", padx=(0, 15))

# Progress bar (in header, after logo)
progress_frame = ctk.CTkFrame(header_frame, fg_color="transparent")

progress_label = ctk.CTkLabel(
    progress_frame, 
    text="Processing...", 
    font=(FONT_FAMILY, FONT_SIZES["small"])
)
progress_label.pack(side="left", padx=(0, 8))

progress_bar = ctk.CTkProgressBar(
    progress_frame, 
    width=300, 
    height=12,
    corner_radius=6,
    progress_color=BUTTON_COLOR
)
progress_bar.pack(side="left")
progress_bar.set(0)
# Don't pack initially - will be shown when needed

# Model selector dropdown (in header, before status)
model_label = ctk.CTkLabel(
    header_frame, 
    text="Model:", 
    font=(FONT_FAMILY, FONT_SIZES["body"])
)
model_label.pack(side="right", padx=(0, 8))

model_dropdown = ctk.CTkOptionMenu(
    header_frame, 
    values=AVAILABLE_MODELS,
    command=on_model_change,
    width=110,
    height=32,
    font=(FONT_FAMILY, FONT_SIZES["body"]),
    corner_radius=6,
    fg_color=BUTTON_COLOR,
    button_color=BUTTON_HOVER_COLOR,
    button_hover_color=BUTTON_COLOR,
    text_color=BUTTON_TEXT_COLOR
)
model_dropdown.set(DEFAULT_MODEL)
model_dropdown.pack(side="right", padx=(0, 15))

if not USE_CUDA:
    model_dropdown.configure(state="disabled")
    model_label.configure(text="Model: (GPU required)")

# License status label (before status label)
license_status_label = ctk.CTkLabel(
    header_frame,
    text="",
    font=(FONT_FAMILY, FONT_SIZES["small"]),
    text_color="#FFAA00"  # Orange for trial, will change to green if unlocked
)
license_status_label.pack(side="right", padx=(0, 15))

# Unlock/Activate button
def show_unlock_dialog():
    """Show the unlock dialog."""
    try:
        from mmvidstoclips_gui_unlock import UnlockDialog
        dialog = UnlockDialog(parent=app)
        dialog.run()
        # Refresh license status after dialog closes
        update_license_status()
    except Exception as e:
        print(f"Error showing unlock dialog: {e}")
        import tkinter.messagebox as messagebox
        messagebox.showerror("Error", f"Could not open unlock dialog: {e}")

unlock_btn = ctk.CTkButton(
    header_frame,
    text="🔓 Unlock",
    width=100,
    height=28,
    font=(FONT_FAMILY, FONT_SIZES["small"]),
    corner_radius=6,
    fg_color="#CC4444",  # Red color matching close button
    hover_color="#FF5555",  # Lighter red on hover
    text_color="white",
    command=show_unlock_dialog
)
unlock_btn.pack(side="right", padx=(0, 10))

def update_license_status():
    """Update the license status display."""
    if not LICENSE_AVAILABLE:
        license_status_label.configure(text="")
        unlock_btn.pack_forget()
        return
    
    try:
        # Initialize trial if needed
        initialize_trial()
        
        status = get_license_status()
        
        if status['unlocked']:
            license_status_label.configure(
                text="✓ UNLOCKED",
                text_color="#44FF44"  # Green
            )
            unlock_btn.configure(
                text="✓ Activated",
                fg_color="#44FF44",
                hover_color="#66FF66",
                state="normal"
            )
        else:
            remaining = status['trial_remaining']
            if remaining > 0:
                license_status_label.configure(
                    text=f"Trial: {remaining} day(s) remaining",
                    text_color="#3c847b"  # Custom color
                )
                unlock_btn.configure(
                    text="🔓 Unlock",
                    fg_color="#CC4444",  # Red color matching close button
                    hover_color="#FF5555",  # Lighter red on hover
                    state="normal"
                )
            elif remaining == 0:
                license_status_label.configure(
                    text="⚠ Trial Expired",
                    text_color="#FF4444"  # Red
                )
                unlock_btn.configure(
                    text="🔓 Activate",
                    fg_color="#CC4444",  # Red color matching close button
                    hover_color="#FF5555",  # Lighter red on hover
                    state="normal"
                )
            else:
                license_status_label.configure(
                    text="Trial Active",
                    text_color="#3c847b"  # Custom color
                )
                unlock_btn.configure(
                    text="🔓 Unlock",
                    fg_color="#CC4444",  # Red color matching close button
                    hover_color="#FF5555",  # Lighter red on hover
                    state="normal"
                )
    except Exception as e:
        print(f"Error updating license status: {e}")
        license_status_label.configure(text="")

# Status label showing GPU/CPU mode
status_label = ctk.CTkLabel(
    header_frame, 
    text="", 
    font=(FONT_FAMILY, FONT_SIZES["body"], "bold")
)
status_label.pack(side="right", padx=15)

# Row 1: Saved Transcripts dropdown and Model dropdown
controls_frame = ctk.CTkFrame(app, fg_color="transparent")
controls_frame.pack(fill="x", padx=15, pady=(10, 10))

# Saved Transcripts label
saved_label = ctk.CTkLabel(
    controls_frame, 
    text="Saved Transcripts:", 
    font=(FONT_FAMILY, FONT_SIZES["body"])
)
saved_label.pack(side="left", padx=(0, 8))

def on_saved_transcript_select(choice):
    """Handle selection of saved transcript - just keeps selection visible, doesn't load."""
    # Selection is now just for display - user must click Load button to actually load
    pass

def load_selected_transcript():
    """Load the currently selected saved transcript into first empty panel, searching right to left."""
    current = saved_transcripts_dropdown.get()
    if not current or current in ["-- Select --", "(No saved transcripts)"]:
        return
    
    result = load_transcript_from_app(current)
    if isinstance(result, tuple):
        text, formatting_tags = result
    else:
        text, formatting_tags = result, None
    
    if text and panels:
        # Search from right to left (reverse order) for first empty panel
        for panel in reversed(panels):
            panel_text = panel.get_text()
            if not panel_text or not panel_text.strip():
                # Found an empty panel - load the transcript here
                panel.set_text(text, formatting_tags=formatting_tags)
                # Associate this panel with the saved file for auto-save
                panel.associated_saved_file = current
                saved_transcripts_dropdown.set("-- Select --")
                return
        
        # If no empty panel found, use the rightmost panel (last in list)
        panels[-1].set_text(text, apply_bold=apply_bold, formatting_tags=formatting_tags)
        # Associate this panel with the saved file for auto-save
        panels[-1].associated_saved_file = current
        saved_transcripts_dropdown.set("-- Select --")

# Get initial list of saved transcripts
initial_saved = get_saved_transcripts()
saved_values = ["-- Select --"] + initial_saved if initial_saved else ["(No saved transcripts)"]

saved_transcripts_dropdown = ctk.CTkOptionMenu(
    controls_frame, 
    values=saved_values,
    command=on_saved_transcript_select,
    width=200,
    height=32,
    font=(FONT_FAMILY, FONT_SIZES["body"]),
    corner_radius=6,
    fg_color=BUTTON_COLOR,
    button_color=BUTTON_HOVER_COLOR,
    button_hover_color=BUTTON_COLOR,
    text_color=BUTTON_TEXT_COLOR
)
saved_transcripts_dropdown.set("-- Select --" if initial_saved else "(No saved transcripts)")
saved_transcripts_dropdown.pack(side="left", padx=(0, 10))

# Load selected transcript button
load_saved_btn = ctk.CTkButton(
    controls_frame,
    text="Load",
    width=60,
    height=32,
    font=(FONT_FAMILY, FONT_SIZES["body"]),
    corner_radius=6,
    fg_color=BUTTON_COLOR,
    hover_color=BUTTON_HOVER_COLOR,
    text_color=BUTTON_TEXT_COLOR,
    command=load_selected_transcript
)
load_saved_btn.pack(side="left", padx=(0, 5))

# Delete saved transcript button
def delete_selected_transcript():
    """Delete the currently selected saved transcript."""
    current = saved_transcripts_dropdown.get()
    if current and current not in ["-- Select --", "(No saved transcripts)"]:
        if delete_saved_transcript(current):
            refresh_saved_transcripts_dropdown()
            saved_transcripts_dropdown.set("-- Select --")

delete_saved_btn = ctk.CTkButton(
    controls_frame,
    text="🗑️",
    width=32,
    height=32,
    font=(FONT_FAMILY, 14),
    corner_radius=6,
    fg_color="#666666",
    hover_color="#888888",
    text_color="white",
    command=delete_selected_transcript
)
delete_saved_btn.pack(side="left", padx=(0, 30))

# Saved Audio section
saved_audio_label = ctk.CTkLabel(
    controls_frame, 
    text="Audio:", 
    font=(FONT_FAMILY, FONT_SIZES["body"])
)
saved_audio_label.pack(side="left", padx=(0, 8))

# Saved Audio Management Functions
SAVED_AUDIO_DIR = os.path.join(APP_DIR, "savedaudio")

# Create saved audio folder if it doesn't exist
if not os.path.exists(SAVED_AUDIO_DIR):
    os.makedirs(SAVED_AUDIO_DIR)

def get_saved_audio_files():
    """Get list of saved audio filenames."""
    if not os.path.exists(SAVED_AUDIO_DIR):
        return []
    files = [f for f in os.listdir(SAVED_AUDIO_DIR) if f.lower().endswith(('.mp3', '.wav'))]
    return sorted(files, key=lambda x: os.path.getmtime(os.path.join(SAVED_AUDIO_DIR, x)), reverse=True)

def save_audio_file(file_path):
    """Copy audio file to savedaudio folder."""
    if not os.path.exists(file_path):
        return None
    
    filename = os.path.basename(file_path)
    dest_path = os.path.join(SAVED_AUDIO_DIR, filename)
    
    # If file already exists, add a number suffix
    counter = 1
    base_name, ext = os.path.splitext(filename)
    while os.path.exists(dest_path):
        new_filename = f"{base_name}_{counter}{ext}"
        dest_path = os.path.join(SAVED_AUDIO_DIR, new_filename)
        counter += 1
    
    try:
        shutil.copy2(file_path, dest_path)
        return os.path.basename(dest_path)
    except Exception as e:
        print(f"Error saving audio file: {e}")
        return None

def delete_saved_audio(filename):
    """Delete a saved audio file."""
    filepath = os.path.join(SAVED_AUDIO_DIR, filename)
    try:
        if os.path.exists(filepath):
            # Try to delete the file
            os.remove(filepath)
            # Verify it was deleted
            if not os.path.exists(filepath):
                return True
            else:
                print(f"Warning: File still exists after deletion attempt: {filepath}")
                return False
    except PermissionError as e:
        print(f"Permission error deleting audio file: {e}")
        print(f"File may be in use by another process: {filepath}")
        return False
    except Exception as e:
        print(f"Error deleting audio file: {e}")
        print(f"File path: {filepath}")
        return False
    return False

def refresh_saved_audio_dropdown():
    """Refresh the saved audio dropdown with current files."""
    files = get_saved_audio_files()
    if files:
        saved_audio_dropdown.configure(values=["-- Select --"] + files)
    else:
        saved_audio_dropdown.configure(values=["(No saved audio)"])

def on_saved_audio_select(choice):
    """Handle selection of saved audio - just keeps selection visible."""
    pass

def load_audio_file():
    """Load an audio file and save it to savedaudio folder."""
    file_path = filedialog.askopenfilename(
        title="Select Audio File to Save",
        filetypes=[("Audio Files", "*.mp3;*.wav"), ("MP3 Files", "*.mp3"), ("WAV Files", "*.wav"), ("All Files", "*.*")]
    )
    if file_path:
        saved_name = save_audio_file(file_path)
        if saved_name:
            refresh_saved_audio_dropdown()
            saved_audio_dropdown.set(saved_name)

# Audio Player Widget
audio_player_frame = None
audio_player_thread = None
audio_player_running = False
audio_player_paused = False
audio_player_position = 0
audio_player_duration = 0
current_audio_file = None              # May be a temp speed-adjusted file
original_audio_file = None             # Always the original selected file
audio_playback_speed = 1.0            # Default playback speed
updating_timeline = False             # Flag to prevent seek during automatic updates
audio_start_time = 0                  # Track when playback started for position calculation
audio_pause_start = 0                 # Track pause time
total_paused_time = 0                 # Total time spent paused
is_seeking = False                    # Flag to prevent position updates during seek
user_dragging_timeline = False        # Track if user is actively dragging the timeline
last_programmatic_slider_value = -1   # Track last programmatically set value
temp_speed_file = None                # Temporary file for speed-adjusted audio


def _build_atempo_filters(speed: float):
    """Split speed into factors between 0.5 and 2.0 for ffmpeg atempo."""
    factors = []
    remaining = speed
    # Handle speeds > 2 by chaining 2.0
    while remaining > 2.0:
        factors.append(2.0)
        remaining /= 2.0
    # Handle speeds < 0.5 by chaining 0.5
    while remaining < 0.5:
        factors.append(0.5)
        remaining /= 0.5
    factors.append(remaining)
    return factors


def create_speed_adjusted_file_ffmpeg(src_path: str, speed: float):
    """
    Create a pitch-preserving speed-adjusted WAV using ffmpeg atempo.
    Returns path to temp wav or None if ffmpeg unavailable or speed==1.
    """
    if speed == 1.0:
        return None
    ffmpeg_path = shutil.which("ffmpeg")
    if not ffmpeg_path:
        return None
    try:
        temp_dir = os.path.join(APP_DIR, "temp_audio")
        try:
            os.makedirs(temp_dir, exist_ok=True)
        except Exception as e:
            print(f"Error creating temp_audio directory in APP_DIR: {e}")
            # Fallback to system temp directory
            import tempfile as tf
            temp_dir = os.path.join(tf.gettempdir(), "MMTranscriptEditor_temp_audio")
            try:
                os.makedirs(temp_dir, exist_ok=True)
            except Exception as e2:
                print(f"Error creating temp_audio in system temp: {e2}")
                raise
        # unique temp file
        fd, out_path = tempfile.mkstemp(prefix=f"atempo_{speed}_", suffix=".wav", dir=temp_dir)
        os.close(fd)
        # build filter chain
        factors = _build_atempo_filters(speed)
        filter_chain = ",".join([f"atempo={f}" for f in factors])
        cmd = [
            ffmpeg_path,
            "-y",
            "-i", src_path,
            "-filter:a", filter_chain,
            "-acodec", "pcm_s16le",
            out_path,
        ]
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if result.returncode != 0 or not os.path.exists(out_path):
            print(f"ffmpeg atempo failed (code {result.returncode}): {result.stderr.decode(errors='ignore')}")
            try:
                os.remove(out_path)
            except:
                pass
            return None
        return out_path
    except Exception as e:
        print(f"Error creating atempo file: {e}")
        return None

def format_time(seconds):
    """Format seconds as MM:SS."""
    if seconds < 0:
        seconds = 0
    mins = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{mins:02d}:{secs:02d}"

def create_audio_player():
    """Create the audio player widget."""
    global audio_player_frame
    
    # Remove existing player if it exists
    if audio_player_frame:
        try:
            audio_player_frame.destroy()
        except:
            pass
    
    # Create player frame (pack it in a new row below controls_frame)
    audio_player_frame = ctk.CTkFrame(app, fg_color=get_theme_color("textbox_bg"), corner_radius=8, height=80)
    # Pack it right after controls_frame
    audio_player_frame.pack(fill="x", padx=15, pady=(0, 10), after=controls_frame)
    
    # Player controls row
    controls_row = ctk.CTkFrame(audio_player_frame, fg_color="transparent")
    controls_row.pack(fill="x", padx=10, pady=10)
    
    # Play/Pause button
    play_pause_btn = ctk.CTkButton(
        controls_row,
        text="▶️",
        width=40,
        height=32,
        font=(FONT_FAMILY, 14),
        corner_radius=6,
        fg_color=BUTTON_COLOR,
        hover_color=BUTTON_HOVER_COLOR,
        text_color=BUTTON_TEXT_COLOR,
        command=toggle_play_pause
    )
    play_pause_btn.pack(side="left", padx=(0, 5))
    
    # Stop button
    stop_btn = ctk.CTkButton(
        controls_row,
        text="⏹️",
        width=40,
        height=32,
        font=(FONT_FAMILY, 14),
        corner_radius=6,
        fg_color=BUTTON_COLOR,
        hover_color=BUTTON_HOVER_COLOR,
        text_color=BUTTON_TEXT_COLOR,
        command=stop_audio
    )
    stop_btn.pack(side="left", padx=(0, 10))
    
    # Time label (current time / total time)
    time_label = ctk.CTkLabel(
        controls_row,
        text="00:00 / 00:00",
        font=(FONT_FAMILY, FONT_SIZES["small"]),
        width=100
    )
    time_label.pack(side="right", padx=(10, 0))
    
    # Close button
    close_btn = ctk.CTkButton(
        controls_row,
        text="✕",
        width=28,
        height=28,
        font=(FONT_FAMILY, 12),
        corner_radius=4,
        fg_color="#666666",
        hover_color="#888888",
        text_color="white",
        command=hide_audio_player
    )
    close_btn.pack(side="right", padx=(5, 0))
    
    # Timeline slider
    timeline_slider = ctk.CTkSlider(
        controls_row,
        from_=0,
        to=100,
        width=400,
        height=20,
        button_color="#CC4444",  # Red color matching close button
        button_hover_color="#FF5555",  # Lighter red on hover
        command=on_timeline_change_with_check
    )
    timeline_slider.pack(side="left", fill="x", expand=True, padx=(0, 10))
    
    # Also try to bind to underlying widget events for immediate feedback
    try:
        # CTkSlider wraps a tkinter Scale widget
        underlying_scale = timeline_slider._slider
        underlying_scale.bind("<Button-1>", lambda e: on_timeline_press())
        underlying_scale.bind("<B1-Motion>", lambda e: on_timeline_drag())
        underlying_scale.bind("<ButtonRelease-1>", lambda e: on_timeline_release())
    except (AttributeError, TypeError):
        # Fallback: try alternative attribute names or just use command callback
        try:
            underlying_scale = timeline_slider.slider
            underlying_scale.bind("<Button-1>", lambda e: on_timeline_press())
            underlying_scale.bind("<B1-Motion>", lambda e: on_timeline_drag())
            underlying_scale.bind("<ButtonRelease-1>", lambda e: on_timeline_release())
        except (AttributeError, TypeError):
            # Command callback will handle it
            pass
    
    # Speed control dropdown
    speed_label = ctk.CTkLabel(controls_row, text="Speed:", font=(FONT_FAMILY, FONT_SIZES["small"]))
    speed_label.pack(side="left", padx=(10, 5))
    
    speed_options = ["0.5x", "0.75x", "1.0x", "1.25x", "1.5x", "1.75x", "2.0x"]
    speed_dropdown = ctk.CTkOptionMenu(
        controls_row,
        values=speed_options,
        command=change_playback_speed,
        width=80,
        height=28,
        font=(FONT_FAMILY, FONT_SIZES["small"]),
        corner_radius=4,
        fg_color=BUTTON_COLOR,
        button_color=BUTTON_HOVER_COLOR,
        button_hover_color=BUTTON_COLOR,
        text_color=BUTTON_TEXT_COLOR
    )
    speed_dropdown.set("1.0x")
    speed_dropdown.pack(side="left", padx=(0, 10))
    
    # Store references for updates
    audio_player_frame.play_pause_btn = play_pause_btn
    audio_player_frame.stop_btn = stop_btn
    audio_player_frame.time_label = time_label
    audio_player_frame.timeline_slider = timeline_slider
    audio_player_frame.speed_dropdown = speed_dropdown
    
    return audio_player_frame

def change_playback_speed(speed_str):
    """Change audio playback speed."""
    global audio_playback_speed, audio_segment, audio_player_position, temp_speed_file, audio_player_duration
    global original_audio_file
    try:
        # Parse speed (e.g., "1.5x" -> 1.5)
        speed = float(speed_str.rstrip('x'))
        audio_playback_speed = speed
        
        # If audio is playing, restart with new speed
        if audio_player_running and original_audio_file:
            was_paused = audio_player_paused
            # Stop current playback but preserve original file
            stop_audio(reset_original=False)
            # Restart with new speed
            play_audio_file(original_audio_file, is_original=True)
            if was_paused:
                toggle_play_pause()  # pause again
    except Exception as e:
        print(f"Error parsing speed: {e}")

def toggle_play_pause():
    """Toggle play/pause state."""
    global audio_player_paused, audio_pause_start, total_paused_time, audio_start_time
    
    if not audio_player_frame or not current_audio_file or not PYGAME_AVAILABLE:
        return
    
    if not audio_player_running:
        # If not running, start playing
        play_audio_file(original_audio_file if original_audio_file else current_audio_file, is_original=True)
        return
    
    if audio_player_paused:
        # Resume playback
        if PYGAME_AVAILABLE:
            try:
                pygame.mixer.music.unpause()
                # Add the paused time to total paused time
                total_paused_time += time.time() - audio_pause_start
                audio_player_paused = False
                audio_player_frame.play_pause_btn.configure(text="⏸️")
            except Exception as e:
                print(f"Error resuming audio: {e}")
    else:
        # Pause playback
        if PYGAME_AVAILABLE:
            try:
                pygame.mixer.music.pause()
                audio_pause_start = time.time()
                audio_player_paused = True
                audio_player_frame.play_pause_btn.configure(text="▶️")
            except Exception as e:
                print(f"Error pausing audio: {e}")

def stop_audio(reset_original=True):
    """Stop audio playback.
    
    Args:
        reset_original: If True, reset original_audio_file. Set to False when restarting playback.
    """
    global audio_player_running, audio_player_paused, audio_player_position, current_audio_file
    global audio_start_time, total_paused_time, audio_pause_start, temp_speed_file, original_audio_file
    
    audio_player_running = False
    audio_player_paused = False
    audio_player_position = 0
    audio_start_time = 0
    total_paused_time = 0
    audio_pause_start = 0
    if reset_original:
        original_audio_file = None
    
    if PYGAME_AVAILABLE:
        try:
            pygame.mixer.music.stop()
            # unload() is available in pygame 2.0+ - release file handle
            if hasattr(pygame.mixer.music, 'unload'):
                pygame.mixer.music.unload()
        except:
            pass
    
    # Clean up temporary speed file only if resetting
    if reset_original:
        if temp_speed_file and os.path.exists(temp_speed_file):
            try:
                os.remove(temp_speed_file)
            except:
                pass
        temp_speed_file = None
        current_audio_file = None
    
    if audio_player_frame:
        audio_player_frame.play_pause_btn.configure(text="▶️")
        audio_player_frame.timeline_slider.set(0)
        audio_player_frame.time_label.configure(text="00:00 / 00:00")

def hide_audio_player():
    """Hide the audio player widget."""
    global audio_player_frame
    if audio_player_frame:
        stop_audio()
        try:
            audio_player_frame.pack_forget()
        except:
            pass

_timeline_release_scheduled = False

def on_timeline_change_with_check(value):
    """Handle timeline change with check for user interaction."""
    global last_programmatic_slider_value, user_dragging_timeline, _timeline_release_scheduled
    
    # If this value matches what we just set programmatically, ignore it
    if abs(value - last_programmatic_slider_value) < 0.1:
        return
    
    # This is a user interaction - handle it
    if not updating_timeline:
        user_dragging_timeline = True
        on_timeline_drag()
        # Cancel any previously scheduled release
        if _timeline_release_scheduled:
            try:
                app.after_cancel(_timeline_release_scheduled)
            except:
                pass
        # Schedule release after user stops dragging (300ms delay)
        _timeline_release_scheduled = app.after(300, on_timeline_release)

def on_timeline_press():
    """Handle timeline slider button press - user starts dragging."""
    global user_dragging_timeline
    user_dragging_timeline = True

def on_timeline_drag(event=None):
    """Handle timeline slider drag - update position display but don't seek yet."""
    global audio_player_position, audio_player_duration
    if not audio_player_frame or not current_audio_file or not user_dragging_timeline:
        return
    
    # Get current slider value
    try:
        slider_value = audio_player_frame.timeline_slider.get()
    except:
        return
    
    # Update position display for visual feedback
    if audio_player_duration > 0:
        new_position = (slider_value / 100.0) * audio_player_duration
        audio_player_position = new_position
        
        # Update time label for immediate feedback
        current_time_str = format_time(audio_player_position)
        total_time_str = format_time(audio_player_duration)
        audio_player_frame.time_label.configure(text=f"{current_time_str} / {total_time_str}")

def on_timeline_release():
    """Handle timeline slider button release - perform actual seek."""
    global audio_player_position, is_seeking, audio_start_time, total_paused_time, user_dragging_timeline, _timeline_release_scheduled
    global audio_player_paused, audio_pause_start, audio_player_running, audio_player_duration, current_audio_file
    
    _timeline_release_scheduled = False
    user_dragging_timeline = False
    
    if not audio_player_frame or not current_audio_file or not PYGAME_AVAILABLE:
        return
    
    if not audio_player_running:
        return
    
    # Get current slider value
    slider_value = audio_player_frame.timeline_slider.get()
    
    # Only seek if user is dragging (not during automatic updates)
    if audio_player_duration > 0:
        new_position = (slider_value / 100.0) * audio_player_duration
        is_seeking = True
        
        try:
            # Stop current playback
            was_paused = audio_player_paused
            pygame.mixer.music.stop()
            
            # Reload and play from beginning
            pygame.mixer.music.load(current_audio_file)
            pygame.mixer.music.play()
            
            # Try to seek to position (limited support in pygame)
            # Note: set_pos() only works for some formats and may not be accurate
            try:
                if new_position > 0:
                    pygame.mixer.music.set_pos(new_position)
            except:
                # If set_pos fails, we'll track position manually
                pass
            
            # Update position tracking
            audio_player_position = new_position
            audio_start_time = time.time() - new_position
            total_paused_time = 0
            
            # Restore pause state if it was paused
            if was_paused:
                pygame.mixer.music.pause()
                audio_pause_start = time.time()
            else:
                audio_player_paused = False
            
            is_seeking = False
        except Exception as e:
            print(f"Error seeking audio: {e}")
            is_seeking = False

def update_audio_player():
    """Update audio player UI (time, timeline)."""
    global audio_player_position, audio_player_duration, updating_timeline
    global audio_start_time, total_paused_time, audio_pause_start, audio_playback_speed
    
    if not audio_player_frame or not current_audio_file or is_seeking or user_dragging_timeline:
        if audio_player_running:
            app.after(100, update_audio_player)
        return
    
    if PYGAME_AVAILABLE and audio_player_running:
        # Calculate actual position based on elapsed time
        if not audio_player_paused:
            current_time = time.time()
            elapsed = current_time - audio_start_time - total_paused_time
            # The mixer plays the (possibly speed-adjusted) file at real-time rate.
            # We want position in original-time scale, so scale by playback speed.
            audio_player_position = max(0, elapsed * audio_playback_speed)
        else:
            # While paused, don't update position
            pass
        
        # Check if playback finished
        if not pygame.mixer.music.get_busy() and not audio_player_paused:
            # Check if we've reached the end
            if audio_player_duration > 0 and audio_player_position >= audio_player_duration - 0.5:
                stop_audio()
                return
            # If not at end but not busy, might have stopped - restart update loop
            elif audio_player_duration > 0 and audio_player_position < audio_player_duration:
                # Still playing, continue
                pass
    
    # Update timeline slider (set flag to prevent seek callback)
    # Only update if user is not dragging
    if audio_player_duration > 0 and not user_dragging_timeline:
        progress = (audio_player_position / audio_player_duration) * 100
        updating_timeline = True
        try:
            new_value = min(100, max(0, progress))
            audio_player_frame.timeline_slider.set(new_value)
            last_programmatic_slider_value = new_value
        except:
            pass
        updating_timeline = False
    
    # Update time label
    current_time_str = format_time(audio_player_position)
    total_time_str = format_time(audio_player_duration)
    audio_player_frame.time_label.configure(text=f"{current_time_str} / {total_time_str}")
    
    # Schedule next update
    if audio_player_running:
        app.after(100, update_audio_player)

def play_audio_file(filepath, is_original=False):
    """Play audio file using pygame.
    
    Args:
        filepath: Path to audio file to play
        is_original: If True, this is the original file (not a temp speed-adjusted file)
    """
    global audio_player_running, audio_player_paused, audio_player_position, audio_player_duration, current_audio_file
    global audio_start_time, total_paused_time, audio_pause_start, original_audio_file, temp_speed_file
    global audio_playback_speed
    
    if not PYGAME_AVAILABLE:
        # Show message and fallback to system player
        import tkinter.messagebox as messagebox
        messagebox.showwarning(
            "Pygame Not Available",
            "Pygame is not installed. The in-app audio player requires pygame.\n\n"
            "To install pygame:\n"
            "1. Open a terminal/command prompt\n"
            "2. Navigate to: c:\\Users\\arlyn\\Documents\\MMTranscriptEditor\n"
            "3. Run: grammateus-env\\Scripts\\activate\n"
            "4. Run: pip install pygame mutagen\n\n"
            "Opening file with system default player instead."
        )
        # Fallback to system player
        try:
            system = platform.system()
            if system == "Windows":
                os.startfile(filepath)
            elif system == "Darwin":  # macOS
                subprocess.run(["open", filepath])
            else:  # Linux
                subprocess.run(["xdg-open", filepath])
        except Exception as e:
            print(f"Error playing audio file: {e}")
        return
    
    try:
        # Stop any currently playing audio (but preserve original_audio_file if we're restarting)
        stop_audio(reset_original=False)
        
        # Initialize pygame mixer if not already done
        if not pygame.mixer.get_init():
            try:
                pygame.mixer.pre_init(frequency=44100, size=-16, channels=2, buffer=512)
                pygame.mixer.init()
            except Exception as e:
                print(f"Error pre-initializing mixer: {e}, trying default init...")
                try:
                    pygame.mixer.init()
                except Exception as e2:
                    print(f"Error initializing pygame mixer: {e2}")
                    import tkinter.messagebox as messagebox
                    messagebox.showerror("Audio Error", f"Cannot initialize audio player:\n{e2}\n\nPlease check that audio drivers are installed.")
                    raise Exception(f"Cannot initialize audio: {e2}")
        
        # Store original file path
        if is_original or original_audio_file is None:
            original_audio_file = filepath
        
        # Handle speed adjustment if needed (only for original files, not temp files)
        play_file = filepath
        if is_original and audio_playback_speed != 1.0:
            # Try pitch-preserving ffmpeg atempo first
            try:
                if temp_speed_file and os.path.exists(temp_speed_file):
                    try:
                        os.remove(temp_speed_file)
                    except:
                        pass
                temp_speed_file = create_speed_adjusted_file_ffmpeg(filepath, audio_playback_speed)
                if temp_speed_file:
                    print(f"Using ffmpeg atempo temp file: {temp_speed_file}")
                    play_file = temp_speed_file
                elif PYDUB_AVAILABLE:
                    # Fallback to pydub resample (will change pitch)
                    audio_segment = AudioSegment.from_file(filepath)
                    new_frame_rate = int(audio_segment.frame_rate * audio_playback_speed)
                    audio_segment = audio_segment._spawn(audio_segment.raw_data, overrides={
                        "frame_rate": new_frame_rate
                    }).set_frame_rate(new_frame_rate)
                    
                    temp_dir = os.path.join(APP_DIR, "temp_audio")
                    try:
                        os.makedirs(temp_dir, exist_ok=True)
                    except Exception as e:
                        print(f"Error creating temp_audio directory: {e}")
                        # Fallback to system temp directory
                        import tempfile as tf
                        temp_dir = os.path.join(tf.gettempdir(), "MMTranscriptEditor_temp_audio")
                        os.makedirs(temp_dir, exist_ok=True)
                    temp_speed_file = os.path.join(temp_dir, f"speed_{audio_playback_speed}_{os.path.basename(filepath)}").replace(" ", "_").replace("\\", "_").replace("/", "_")
                    if not temp_speed_file.endswith(".wav"):
                        temp_speed_file += ".wav"
                    
                    print(f"Creating speed-adjusted audio file (pydub): {temp_speed_file}")
                    audio_segment.export(temp_speed_file, format="wav", parameters=["-acodec", "pcm_s16le"])
                    if os.path.exists(temp_speed_file):
                        print(f"Pydub speed-adjusted file created: {os.path.getsize(temp_speed_file)} bytes")
                        play_file = temp_speed_file
                    else:
                        temp_speed_file = None
                else:
                    temp_speed_file = None
            except Exception as e:
                print(f"Error creating speed-adjusted audio: {e}")
                import traceback
                traceback.print_exc()
                play_file = filepath
                temp_speed_file = None
        
        # Get duration first (needed for seeking) - use original file for duration calculation
        audio_player_duration = 0
        duration_file = original_audio_file if original_audio_file else filepath
        if MUTAGEN_AVAILABLE:
            try:
                audio_file = MutagenFile(duration_file)  # Use original file for duration
                if audio_file and hasattr(audio_file, 'info') and hasattr(audio_file.info, 'length'):
                    original_duration = audio_file.info.length
                    # Adjust duration for speed (faster speed = shorter playback time)
                    audio_player_duration = original_duration / audio_playback_speed
            except:
                pass
        
        # Fallback: estimate duration (rough approximation)
        if audio_player_duration == 0:
            try:
                file_size = os.path.getsize(duration_file)  # Use original file for size
                # Rough estimate: 1MB ≈ 1 minute for MP3/WAV
                original_duration = (file_size / (1024 * 1024)) * 60
                # Adjust duration for speed
                audio_player_duration = original_duration / audio_playback_speed
            except:
                audio_player_duration = 0
        
        # Re-init mixer to match audio file properties if possible
        try:
            # Use pydub to probe the file for sample rate / channels
            if PYDUB_AVAILABLE:
                probe_segment = AudioSegment.from_file(play_file)
                pygame.mixer.quit()
                try:
                    pygame.mixer.init(frequency=probe_segment.frame_rate, channels=probe_segment.channels)
                except:
                    # If specific init fails, try default
                    pygame.mixer.init()
            else:
                # If pydub not available, ensure mixer is initialized
                if not pygame.mixer.get_init():
                    pygame.mixer.init()
        except Exception as e:
            print(f"Warning: Could not re-init mixer with file properties: {e}")
            # Fallback to default init if probing fails
            if not pygame.mixer.get_init():
                try:
                    pygame.mixer.init()
                except Exception as e2:
                    print(f"Error initializing mixer: {e2}")
                    raise
        
        # Load and play audio (use speed-adjusted file if created)
        print(f"Loading audio file: {play_file}")
        if not os.path.exists(play_file):
            error_msg = f"Audio file not found: {play_file}"
            print(error_msg)
            import tkinter.messagebox as messagebox
            messagebox.showerror("File Not Found", error_msg)
            raise FileNotFoundError(error_msg)
        
        try:
            pygame.mixer.music.load(play_file)
        except Exception as e:
            error_msg = f"Error loading audio file: {e}"
            print(error_msg)
            import tkinter.messagebox as messagebox
            messagebox.showerror("Audio Error", error_msg)
            raise
        
        try:
            pygame.mixer.music.play()
        except Exception as e:
            error_msg = f"Error starting audio playback: {e}"
            print(error_msg)
            import tkinter.messagebox as messagebox
            messagebox.showerror("Audio Error", error_msg)
            raise
        
        # Verify playback started
        if not pygame.mixer.music.get_busy():
            print("Warning: Audio playback did not start. Retrying...")
            try:
                pygame.mixer.music.play()
            except Exception as e:
                print(f"Retry failed: {e}")
                import tkinter.messagebox as messagebox
                messagebox.showerror("Audio Error", f"Audio playback failed to start:\n{e}")
                raise
        
        current_audio_file = play_file  # Store actual file being played (may be temp file)
        audio_player_running = True
        audio_player_paused = False
        audio_player_position = 0
        audio_start_time = time.time()
        total_paused_time = 0
        audio_pause_start = 0
        
        print(f"Audio playback started. Speed: {audio_playback_speed}x, File: {play_file}")
        
        # Update play button to show pause
        if audio_player_frame:
            audio_player_frame.play_pause_btn.configure(text="⏸️")
        
        # Start UI updates
        update_audio_player()
        
    except Exception as e:
        print(f"Error playing audio: {e}")

def play_selected_audio():
    """Play the currently selected saved audio file."""
    current = saved_audio_dropdown.get()
    if not current or current in ["-- Select --", "(No saved audio)"]:
        return
    
    filepath = os.path.join(SAVED_AUDIO_DIR, current)
    if not os.path.exists(filepath):
        return
    
    # Create/show player widget
    create_audio_player()
    
    # Play the audio
    play_audio_file(filepath, is_original=True)

def delete_selected_audio():
    """Delete the currently selected saved audio file."""
    global current_audio_file, audio_player_running, audio_player_frame
    
    import tkinter.messagebox as messagebox
    import time
    
    current = saved_audio_dropdown.get()
    if not current or current in ["-- Select --", "(No saved audio)"]:
        messagebox.showwarning("No Selection", "Please select an audio file to delete.")
        return
    
    # Check if the file is currently playing
    filepath = os.path.join(SAVED_AUDIO_DIR, current)
    is_currently_playing = (current_audio_file == filepath and audio_player_running)
    
    # Stop playback if this file is playing
    if is_currently_playing:
        stop_audio()
        if audio_player_frame:
            hide_audio_player()
        # Clear the current audio file reference
        current_audio_file = None
        # Give Windows time to release the file handle
        app.update()  # Process any pending events
        time.sleep(0.2)  # Small delay to allow file handle release
    
    # Ask for confirmation
    response = messagebox.askyesno(
        "Delete Audio File",
        f"Are you sure you want to delete '{current}'?",
        icon="question"
    )
    
    if response:
        # Ensure pygame has released the file if it was loaded (even if not currently playing)
        if PYGAME_AVAILABLE:
            try:
                # Always stop and unload to ensure file is released
                pygame.mixer.music.stop()
                # unload() is available in pygame 2.0+
                if hasattr(pygame.mixer.music, 'unload'):
                    pygame.mixer.music.unload()
                current_audio_file = None
                app.update()  # Process events
                time.sleep(0.3)  # Give Windows time to release the file handle
            except Exception as e:
                print(f"Error releasing audio file: {e}")
        
        # Try to delete the file
        result = delete_saved_audio(current)
        if result:
            refresh_saved_audio_dropdown()
            saved_audio_dropdown.set("-- Select --")
            messagebox.showinfo("Success", f"'{current}' has been deleted.")
        else:
            # Try one more time after a longer delay
            app.update()
            time.sleep(0.5)
            result = delete_saved_audio(current)
            if result:
                refresh_saved_audio_dropdown()
                saved_audio_dropdown.set("-- Select --")
                messagebox.showinfo("Success", f"'{current}' has been deleted.")
            else:
                error_msg = f"Failed to delete '{current}'.\n\n"
                error_msg += "The file may be in use by another program or locked.\n\n"
                error_msg += "Please try:\n"
                error_msg += "1. Close any programs that might be using this file\n"
                error_msg += "2. Wait a few seconds and try again\n"
                error_msg += f"3. Manually delete the file from:\n{SAVED_AUDIO_DIR}"
                messagebox.showerror("Error", error_msg)

# Get initial list of saved audio files
initial_audio = get_saved_audio_files()
audio_values = ["-- Select --"] + initial_audio if initial_audio else ["(No saved audio)"]

saved_audio_dropdown = ctk.CTkOptionMenu(
    controls_frame, 
    values=audio_values,
    command=on_saved_audio_select,
    width=200,
    height=32,
    font=(FONT_FAMILY, FONT_SIZES["body"]),
    corner_radius=6,
    fg_color=BUTTON_COLOR,
    button_color=BUTTON_HOVER_COLOR,
    button_hover_color=BUTTON_COLOR,
    text_color=BUTTON_TEXT_COLOR
)
saved_audio_dropdown.set("-- Select --" if initial_audio else "(No saved audio)")
saved_audio_dropdown.pack(side="left", padx=(0, 10))

# Load audio file button (to save new files)
load_audio_btn = ctk.CTkButton(
    controls_frame,
    text="📁",
    width=32,
    height=32,
    font=(FONT_FAMILY, 14),
    corner_radius=6,
    fg_color=BUTTON_COLOR,
    hover_color=BUTTON_HOVER_COLOR,
    text_color=BUTTON_TEXT_COLOR,
    command=load_audio_file
)
load_audio_btn.pack(side="left", padx=(0, 5))

# Play selected audio button
play_audio_btn = ctk.CTkButton(
    controls_frame,
    text="▶️",
    width=32,
    height=32,
    font=(FONT_FAMILY, 14),
    corner_radius=6,
    fg_color=BUTTON_COLOR,
    hover_color=BUTTON_HOVER_COLOR,
    text_color=BUTTON_TEXT_COLOR,
    command=play_selected_audio
)
play_audio_btn.pack(side="left", padx=(0, 5))

# Delete saved audio button
delete_audio_btn = ctk.CTkButton(
    controls_frame,
    text="🗑️",
    width=32,
    height=32,
    font=(FONT_FAMILY, 14),
    corner_radius=6,
    fg_color="#666666",
    hover_color="#888888",
    text_color="white",
    command=delete_selected_audio
)
delete_audio_btn.pack(side="left", padx=(0, 30))

# Theme toggle button (right side of controls)
def toggle_theme():
    """Toggle between Light and Dark themes."""
    global CURRENT_THEME
    CURRENT_THEME = "Light" if CURRENT_THEME == "Dark" else "Dark"
    ctk.set_appearance_mode(CURRENT_THEME)
    
    # Update theme button text
    theme_btn.configure(text="🌙" if CURRENT_THEME == "Light" else "☀️")
    
    # Update main window background color
    apply_main_bg_color()
    
    # Update all panel colors
    for panel in panels:
        panel.textbox.configure(
            fg_color=get_theme_color("textbox_bg"),
            border_color=get_theme_color("textbox_border"),
            text_color=get_theme_color("textbox_text")
        )
        panel.resize_handle.configure(fg_color=get_theme_color("resize_handle"))
        panel.mini_bar.configure(fg_color=get_theme_color("mini_bar"))
        panel.preview_frame.configure(fg_color=get_theme_color("textbox_bg"))
    
    # Save theme preference
    user_settings["theme"] = CURRENT_THEME
    save_settings(user_settings)

theme_btn = ctk.CTkButton(
    controls_frame,
    text="☀️" if CURRENT_THEME == "Dark" else "🌙",
    width=36,
    height=32,
    font=(FONT_FAMILY, 16),
    corner_radius=6,
    fg_color=BUTTON_COLOR,
    hover_color=BUTTON_HOVER_COLOR,
    text_color=BUTTON_TEXT_COLOR,
    command=toggle_theme
)
theme_btn.pack(side="right", padx=(10, 0))

# Theme label
theme_label = ctk.CTkLabel(
    controls_frame,
    text="Theme:",
    font=(FONT_FAMILY, FONT_SIZES["body"])
)
theme_label.pack(side="right", padx=(0, 5))

# Check for Updates button
def on_check_for_updates():
    """Handle check for updates button click."""
    import tkinter.messagebox as messagebox
    
    # Disable button temporarily to prevent multiple clicks
    update_btn.configure(state="disabled", text="Checking...")
    app.update()
    
    try:
        success, message = check_for_update()
        
        if success:
            # Check if message contains download link
            if "Would you like to open the download page?" in message:
                # Ask if user wants to open download page
                response = messagebox.askyesno("Update Available", message)
                if response:
                    # Extract URL from message or use releases page
                    releases_url = "https://github.com/arlynks/MMTranscriptEditor/releases/latest"
                    try:
                        import webbrowser
                        webbrowser.open(releases_url)
                    except Exception:
                        messagebox.showinfo("Download", f"Please visit: {releases_url}")
            else:
                # Up to date
                messagebox.showinfo("Update Check", message)
        else:
            # Error occurred
            messagebox.showerror("Update Check", message)
    except Exception as e:
        messagebox.showerror("Update Check", "Unable to check for updates right now. Please try again later.")
    finally:
        # Re-enable button
        update_btn.configure(state="normal", text="🔍 Check Updates")

update_btn = ctk.CTkButton(
    controls_frame,
    text="🔍 Check Updates",
    width=120,
    height=32,
    font=(FONT_FAMILY, FONT_SIZES["body"]),
    corner_radius=6,
    fg_color=BUTTON_COLOR,
    hover_color=BUTTON_HOVER_COLOR,
    text_color=BUTTON_TEXT_COLOR,
    command=on_check_for_updates
)
update_btn.pack(side="right", padx=(10, 0))

# Row 2: Plus button and Panels container
panels_row = ctk.CTkFrame(app, fg_color="transparent")
panels_row.pack(fill="both", expand=True, padx=10, pady=10)

# Plus button on far left
plus_frame = ctk.CTkFrame(panels_row, fg_color="transparent", width=50)
plus_frame.pack(side="left", fill="y", padx=(0, 5))

plus_button = ctk.CTkButton(
    plus_frame,
    text="+",
    width=40,
    height=40,
    font=(FONT_FAMILY, 24, "bold"),
    corner_radius=8,
    fg_color=BUTTON_COLOR,
    hover_color=BUTTON_HOVER_COLOR,
    text_color=BUTTON_TEXT_COLOR,
    command=add_new_panel
)
plus_button.pack(pady=10)

# Container for all transcript panels (horizontal scrollable area)
panels_container = ctk.CTkFrame(panels_row, fg_color="transparent")
panels_container.pack(side="left", fill="both", expand=True)

# Initialize default panels (two Transcript panels on startup)
transcript_panel = create_panel()
transcript_panel_2 = create_panel()

# --- Functions ---

def split_into_sentences(text):
    """
    Split text into sentences using a simple, reliable character-by-character approach.
    Returns a list of sentences.
    """
    if not text:
        return []
    
    try:
        sentences = []
        current_sentence = []
        i = 0
        
        # Common abbreviations to skip
        abbrevs = ['Mr.', 'Mrs.', 'Ms.', 'Dr.', 'Prof.', 'Sr.', 'Jr.', 'Rev.', 
                   'Gen.', 'Col.', 'Lt.', 'Sgt.', 'Capt.', 'St.', 'Ave.', 'Rd.', 
                   'Blvd.', 'etc.', 'i.e.', 'e.g.', 'vs.']
        
        while i < len(text):
            char = text[i]
            current_sentence.append(char)
            
            # Check if we hit a sentence boundary
            if char in '.!?':
                # Look ahead to check if this is really the end of a sentence
                rest_of_text = text[i:]
                
                # Check if this is part of an abbreviation
                is_abbrev = False
                for abbrev in abbrevs:
                    if rest_of_text.startswith('.' ) and i > 0:
                        # Check previous characters for abbreviation
                        start_pos = max(0, i - 10)
                        context = text[start_pos:i+1]
                        if any(abbr.rstrip('.') in context for abbr in abbrevs):
                            is_abbrev = True
                            break
                
                # Check if it's a decimal number (e.g., 3.16)
                is_decimal = False
                if char == '.' and i > 0 and i < len(text) - 1:
                    if text[i-1].isdigit() and text[i+1].isdigit():
                        is_decimal = True
                
                # If not an abbreviation or decimal, and followed by space or end of text
                if not is_abbrev and not is_decimal:
                    # Check what comes next
                    next_chars = text[i+1:i+3] if i+1 < len(text) else ""
                    
                    # This is a sentence boundary if:
                    # - We're at the end of text, OR
                    # - Next char is whitespace (space, newline, etc.)
                    if i == len(text) - 1 or (next_chars and next_chars[0] in ' \n\t\r'):
                        # Save this sentence
                        sentence = ''.join(current_sentence).strip()
                        if len(sentence) > 3:  # Only keep sentences with substance
                            sentences.append(sentence)
                        current_sentence = []
                        # Skip whitespace after sentence
                        i += 1
                        while i < len(text) and text[i] in ' \n\t\r':
                            i += 1
                        continue
            
            i += 1
        
        # Add any remaining text as a sentence
        if current_sentence:
            sentence = ''.join(current_sentence).strip()
            if len(sentence) > 3:
                sentences.append(sentence)
        
        # If we didn't find any sentences, return the whole text as one sentence
        return sentences if sentences else [text.strip()]
        
    except Exception as e:
        print(f"Sentence split error: {e}")
        # Ultimate fallback: just split on period-space
        return [s.strip() + '.' for s in text.split('. ') if s.strip()]

def format_transcript(text):
    """
    Format transcript into readable paragraph blocks.
    Splits text on sentence-ending punctuation (.!?) and groups 3-4 sentences per paragraph.
    Returns the result joined with double newlines for clean breaks in the UI.
    """
    if not text or not text.strip():
        return text
    
    try:
        # Clean up the text first - normalize whitespace
        text = ' '.join(text.split())
        
        # Split into sentences using sentence-ending punctuation (.!?)
        sentences = split_into_sentences(text)
        
        if not sentences:
            return text
        
        # Group 3-4 sentences per paragraph
        paragraphs = []
        chunk_size = 4  # Target 4 sentences per paragraph
        
        i = 0
        while i < len(sentences):
            remaining = len(sentences) - i
            
            # If 4 or fewer sentences left, put them all in one paragraph
            if remaining <= 4:
                chunk = sentences[i:]
                i = len(sentences)
            else:
                # Take 3-4 sentences (prefer 4, but can use 3)
                # Use 4 if we have enough, otherwise use what's left
                chunk = sentences[i:i + chunk_size]
                i += chunk_size
            
            # Join sentences into a paragraph
            paragraph = " ".join(chunk)
            
            # Ensure proper ending punctuation
            if paragraph and paragraph[-1] not in '.!?':
                paragraph += "."
            
            if paragraph.strip():
                paragraphs.append(paragraph)
        
        # Join paragraphs with double newlines for visual separation
        return "\n\n".join(paragraphs)
        
    except Exception as e:
        print(f"Format transcript error: {e}")
        # Return original text if formatting fails
        return text


# Summarize text using Hugging Face BART model
def summarize_text(text, progress_callback=None):
    """Generate a summary of the transcript text."""
    # Remove extra whitespace and newlines for processing
    clean_text = " ".join(text.split())
    
    # BART has a max input of ~1024 tokens, so chunk by ~800 words to be safe
    words = clean_text.split()
    chunk_size = 800
    chunks = []
    
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i + chunk_size])
        if len(chunk.split()) >= 30:  # Only include chunks worth summarizing
            chunks.append(chunk)
    
    if not chunks:
        return "Text too short to summarize."
    
    # Summarize each chunk with progress updates
    summaries = []
    total_chunks = len(chunks)
    
    # Get summarizer (lazy-loaded)
    summarizer = get_summarizer()
    if summarizer is None:
        return "Error: Unable to load summarization model. Please check your internet connection and try again."
    
    for idx, chunk in enumerate(chunks):
        try:
            result = summarizer(chunk, max_length=150, min_length=30, do_sample=False)
            summaries.append(result[0]["summary_text"])
        except Exception as e:
            summaries.append(f"[Error summarizing chunk: {str(e)}]")
        
        # Update progress (50-100% range for summarization phase)
        if progress_callback:
            progress = 50 + int(((idx + 1) / total_chunks) * 50)
            progress_callback(progress, f"Summarizing... ({idx + 1}/{total_chunks} chunks)")
    
    return "\n\n".join(summaries)

# Initialize status label on startup
update_status_label()

# Initialize license system and check status
if LICENSE_AVAILABLE:
    try:
        # Check if app can run
        if not can_run():
            import tkinter.messagebox as messagebox
            response = messagebox.askyesno(
                "Trial Expired",
                "Your trial period has expired.\n\n"
                "Would you like to enter a license key to unlock the full version?",
                icon="warning"
            )
            if response:
                show_unlock_dialog()
            else:
                app.destroy()
                sys.exit(0)
        
        # Update license status display
        update_license_status()
    except Exception as e:
        print(f"Error initializing license system: {e}")

app.mainloop()
